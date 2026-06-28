#!/usr/bin/env python3
"""Content-level automatic fixes for weekly report Word documents."""
import re

from docx import Document
from docx.oxml.ns import qn


def fix_empty_numbering(doc_path):
    """Remove w:numPr from paragraphs that have no visible text."""
    doc = Document(doc_path)
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        continue
                    pPr = para._p.find(qn('w:pPr'))
                    if pPr is None:
                        continue
                    numPr = pPr.find(qn('w:numPr'))
                    if numPr is not None:
                        pPr.remove(numPr)
                        count += 1
    doc.save(doc_path)
    return f'Removed {count} empty numbering paragraphs'


def fix_cjk_fonts(doc_path, font_name='微软雅黑'):
    """Ensure CJK runs have an eastAsia font set."""
    doc = Document(doc_path)
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        if not re.search(r'[一-鿿]', run.text):
                            continue
                        rpr = run._element.get_or_add_rPr()
                        rFonts = rpr.find(qn('w:rFonts'))
                        if rFonts is None:
                            from docx.oxml import OxmlElement
                            rFonts = OxmlElement('w:rFonts')
                            rpr.append(rFonts)
                        if not rFonts.get(qn('w:eastAsia')):
                            rFonts.set(qn('w:eastAsia'), font_name)
                            count += 1
    doc.save(doc_path)
    return f'Fixed {count} CJK runs missing eastAsia font'


FIX_DISPATCH = {
    '空编号': fix_empty_numbering,
    'eastAsia': fix_cjk_fonts,
}


def apply_fix(doc_path, issue):
    """Apply the first matching automatic fix for an issue string.

    Returns the fix description, or None if no fix matches.
    """
    for keyword, fix_func in FIX_DISPATCH.items():
        if keyword in issue:
            return fix_func(doc_path)
    return None
