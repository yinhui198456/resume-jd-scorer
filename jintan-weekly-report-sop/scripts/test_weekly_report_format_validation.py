#!/usr/bin/env python3
import os
import sys
import unittest
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(__file__))

from report_engine_v9 import ReportGenerator
from validate_report_tone import check_quality_language
from validate_weekly_report import (
    find_empty_numbered_paragraphs,
    find_next_week_plan_issues,
    find_duplicate_cell_content,
)


def add_numbering(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    ppr.append(num_pr)


def excel_date(value):
    return str((value - datetime(1899, 12, 30)).days)


def make_generator():
    config = {
        "project": {},
        "format": {},
        "filter": {},
        "source": {
            "columns": {
                "work_unit": "工作单元",
                "stage": "任务阶段",
                "task": "任务",
                "notes": "备注",
                "progress": "进度",
                "status": "状态",
                "plan_start": "计划开始",
                "plan_end": "计划完成",
            }
        },
    }
    generator = ReportGenerator(config, {})
    generator.monday = datetime(2026, 6, 22)
    generator.friday = datetime(2026, 6, 26)
    generator.sunday = datetime(2026, 6, 28)
    generator.next_monday = datetime(2026, 6, 29)
    generator.next_friday = datetime(2026, 7, 3)
    return generator


class WeeklyReportFormatValidationTests(unittest.TestCase):
    def test_validator_detects_empty_numbered_paragraphs(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        empty_numbered = cell.add_paragraph()
        add_numbering(empty_numbered)
        cell.add_paragraph("有效正文")

        issues = find_empty_numbered_paragraphs(doc)

        self.assertEqual(len(issues), 1)
        self.assertIn("空编号段落", issues[0])

    def test_detects_duplicate_merged_cells(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=4)
        for j in range(4):
            table.cell(0, j).text = "Same duplicate text"

        issues = find_duplicate_cell_content(doc)

        self.assertEqual(len(issues), 1)
        self.assertIn("重复", issues[0])

    def test_generator_clears_empty_numbered_paragraphs_from_target_cell(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        empty_numbered = cell.add_paragraph()
        add_numbering(empty_numbered)
        cell.add_paragraph("旧正文")

        generator = ReportGenerator({"project": {}, "format": {}, "source": {"columns": {}}}, {})
        paragraph = generator._clear_cell_content(cell)

        self.assertEqual(paragraph.text, "")
        self.assertEqual(len(cell.paragraphs), 1)
        pPr = paragraph._p.find(qn("w:pPr"))
        self.assertTrue(
            pPr is None or pPr.find(qn("w:numPr")) is None,
            "cleared cell must not retain numbering",
        )

    def test_tone_validator_flags_low_quality_raw_notes(self):
        text = "用于内部讨论内部讨论；目前稳序尚未完成；乡下反馈他们的数据都翻翻了；smartbi已经通知 smart 开发；先看一下；找一下"

        issues = check_quality_language(text)

        self.assertGreaterEqual(len(issues), 4)
        self.assertTrue(any("重复短语" in issue for issue in issues))
        self.assertTrue(any("疑似错词" in issue for issue in issues))
        self.assertTrue(any("口语化" in issue for issue in issues))
        self.assertTrue(any("术语写法" in issue for issue in issues))
        self.assertTrue(any("看一下" in issue for issue in issues))
        self.assertTrue(any("找一下" in issue for issue in issues))

    def test_tone_validator_flags_action_vague_phrases(self):
        issues = check_quality_language("先看一下火树是否有统计数据；手工数据的台账 SVN 上找一下；逻辑需要再和信息科沟通一下")

        self.assertTrue(any("看一下" in issue for issue in issues))
        self.assertTrue(any("找一下" in issue for issue in issues))
        self.assertTrue(any("一下" in issue for issue in issues))

    def test_next_week_plan_excludes_completed_tasks_and_generates_future_actions(self):
        generator = make_generator()
        plans = [
            {
                "工作单元": "医共体运营管理驾驶舱",
                "任务阶段": "数据应用规划设计",
                "任务": "高保真设计",
                "备注": "已经根据分院反馈意见和总院沟通，完善最终的指标清单",
                "进度": "100%",
                "状态": "进行中",
                "计划开始": excel_date(datetime(2026, 6, 10)),
                "计划完成": excel_date(datetime(2026, 6, 30)),
            },
            {
                "工作单元": "医共体运营管理驾驶舱",
                "任务阶段": "数据应用规划设计",
                "任务": "高保真评审",
                "备注": "需要确认评审的形式及要求",
                "进度": "90%",
                "状态": "进行中",
                "计划开始": excel_date(datetime(2026, 6, 12)),
                "计划完成": excel_date(datetime(2026, 7, 3)),
            },
            {
                "工作单元": "医共体运营管理驾驶舱",
                "任务阶段": "测试与发布",
                "任务": "UAT测试及系统缺陷修复",
                "备注": "调研科室签字",
                "进度": "0%",
                "状态": "未开始",
                "计划开始": excel_date(datetime(2026, 6, 30)),
                "计划完成": excel_date(datetime(2026, 7, 4)),
            },
        ]

        items = generator.build_progress_items(plans, "next_week", all_plans=plans)
        task_texts = [item["text"] for item in items if item["type"] == "task"]
        joined = "\n".join(task_texts)

        self.assertNotIn("高保真设计", joined)
        self.assertNotIn("继续推进", joined)
        self.assertIn("完成高保真评审", joined)
        self.assertIn("评审", joined)
        self.assertIn("启动UAT测试及系统缺陷修复", joined)
        self.assertIn("开展测试验证", joined)

    def test_validator_rejects_status_replay_in_next_week_plan(self):
        doc = Document()
        table = doc.add_table(rows=9, cols=1)
        table.cell(8, 0).text = "高保真设计：已经完成（继续推进）100%"

        issues = find_next_week_plan_issues(doc)

        self.assertTrue(any("继续推进" in issue for issue in issues))
        self.assertTrue(any("100%" in issue for issue in issues))


if __name__ == "__main__":
    unittest.main(verbosity=2)
