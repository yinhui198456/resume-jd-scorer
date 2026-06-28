#!/usr/bin/env python3
import os
import sys
import tempfile
import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(__file__))

from report_auto_fix import fix_empty_numbering, fix_cjk_fonts


class ReportAutoFixTests(unittest.TestCase):
    def test_fix_empty_numbering(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        p = cell.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)

        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, 'test_empty_numbering.docx')
            doc.save(path)
            msg = fix_empty_numbering(path)

            doc2 = Document(path)
            pPr2 = doc2.tables[0].cell(0, 0).paragraphs[0]._p.find(qn('w:pPr'))
            self.assertTrue(
                pPr2 is None or pPr2.find(qn('w:numPr')) is None,
                'empty numbering should be removed',
            )
            self.assertIn('1', msg)

    def test_fix_cjk_fonts(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        run = cell.paragraphs[0].add_run('中文内容')
        # Ensure no eastAsia font is set
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is not None:
            if qn('w:eastAsia') in rFonts.attrib:
                del rFonts.attrib[qn('w:eastAsia')]

        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, 'test_cjk_fonts.docx')
            doc.save(path)
            msg = fix_cjk_fonts(path)

            doc2 = Document(path)
            rpr2 = doc2.tables[0].cell(0, 0).paragraphs[0].runs[0]._element.find(qn('w:rPr'))
            rFonts2 = rpr2.find(qn('w:rFonts'))
            self.assertIsNotNone(rFonts2)
            self.assertEqual(rFonts2.get(qn('w:eastAsia')), '微软雅黑')
            self.assertIn('1', msg)


if __name__ == '__main__':
    unittest.main(verbosity=2)
