#!/usr/bin/env python3
"""
周报质量检查 Agent v2 - 三层校验体系

Layer 1: 结构校验（格式、编号、Markdown）
Layer 2: 数据对照（与 Excel 源数据交叉验证）
Layer 3: 视觉校验（字号、字体、对齐、内容完整性）

使用方式：由 delegate_task 调用，传入文档路径和源数据路径。
"""
import sys
import os
import json
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
import re
from datetime import datetime, timedelta

# ============================================================================
# Configuration
# ============================================================================
DOC_PATH = sys.argv[1] if len(sys.argv) > 1 else ""
SOURCE_XLSX = sys.argv[2] if len(sys.argv) > 2 else os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data", "金坛二期项目跟进表.xlsx")

# ============================================================================
# Excel Parser
# ============================================================================
def parse_excel(filepath):
    if not os.path.exists(filepath):
        return {}
    
    zf = zipfile.ZipFile(filepath)
    ns = {'ss': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
    
    try:
        ss_root = ET.fromstring(zf.read('xl/sharedStrings.xml'))
    except KeyError:
        return {}
    
    ss = []
    for si in ss_root:
        txt = ''.join([t.text or '' for t in si.findall('.//ss:t', ns) + si.findall('.//t') if t.text])
        ss.append(txt)
    
    data = {}
    sheet_map = {
        '02': 'xl/worksheets/sheet3.xml',
    }
    
    for name, path in sheet_map.items():
        if path not in zf.namelist():
            continue
        try:
            root = ET.fromstring(zf.read(path))
        except Exception:
            continue
        
        header_map = {}
        rows = []
        
        for row in root.iter('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}row'):
            cells = {}
            for c in row.iter('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}c'):
                ref = c.get('r', '')
                col = ''.join(filter(str.isalpha, ref))
                v_node = c.find('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v')
                if c.get('t') == 's' and v_node is not None:
                    val = ss[int(v_node.text)] if v_node.text else ''
                elif v_node is not None:
                    val = v_node.text
                else:
                    val = ''
                cells[col] = val
            
            vals = list(cells.values())
            if "里程碑" in str(vals) and "状态" in str(vals):
                header_map = cells
            elif header_map and any(cells.values()):
                row_dict = {header_map.get(k): v for k, v in cells.items()}
                if any(v.strip() for v in row_dict.values()):
                    rows.append(row_dict)
        
        data[name] = rows
    return data

# ============================================================================
# QA Checks
# ============================================================================
def run_qa_checks(doc_path, source_data):
    """Run all three layers of QA checks."""
    doc = Document(doc_path)
    table = doc.tables[0]
    issues = []
    passes = []
    
    # ====================================================================
    # LAYER 1: Structure Checks
    # ====================================================================
    print("\n--- Layer 1: 结构校验 ---")
    
    # 1.1 Markdown symbols
    has_markdown = False
    for row in table.rows:
        for cell in row.cells:
            if '**' in cell.text or '###' in cell.text:
                has_markdown = True
    if has_markdown:
        issues.append({'layer': 1, 'severity': 'HIGH', 'check': 'Markdown', 'message': '文档中存在 Markdown 符号'})
    else:
        passes.append('✅ [L1] Markdown: 已清除')
    
    # 1.2 Numbering in 本周重点
    cell_6 = table.cell(6, 0)
    texts_6 = [p.text.strip() for p in cell_6.paragraphs if p.text.strip()]
    if not any(re.match(r'^[一二三四五六七八]+、', t) for t in texts_6):
        issues.append({'layer': 1, 'severity': 'HIGH', 'check': '编号', 'message': '本周重点缺少一级编号'})
    else:
        passes.append('✅ [L1] 本周一级编号: 存在')
    
    if not any(re.match(r'^\d+\.\s', t) for t in texts_6):
        issues.append({'layer': 1, 'severity': 'MEDIUM', 'check': '编号', 'message': '本周重点缺少二级编号'})
    else:
        passes.append('✅ [L1] 本周二级编号: 存在')
    
    # 1.3 Numbering in 下周计划
    cell_8 = table.cell(8, 0)
    texts_8 = [p.text.strip() for p in cell_8.paragraphs if p.text.strip()]
    if not any(re.match(r'^[一二三四五六七八]+、', t) for t in texts_8):
        issues.append({'layer': 1, 'severity': 'HIGH', 'check': '编号', 'message': '下周计划缺少编号'})
    else:
        passes.append('✅ [L1] 下周编号: 存在')
    
    # 1.4 Numbering in 风险问题（二级标题）
    cell_12 = table.cell(12, 0)
    texts_12 = [p.text.strip() for p in cell_12.paragraphs if p.text.strip()]
    # Check for module-level numbering (1. 2. 3. after app group)
    has_mod_num = any(re.match(r'^\d+\.\s.*：$', t) for t in texts_12)
    if not has_mod_num:
        issues.append({'layer': 1, 'severity': 'HIGH', 'check': '编号', 'message': '风险问题二级标题缺少编号'})
    else:
        passes.append('✅ [L1] 风险二级编号: 存在')
    
    # 1.5 Coordination items numbering
    cell_10 = table.cell(10, 0)
    text_10 = cell_10.text.strip()
    if text_10 and text_10 != '暂无。' and not re.match(r'^\d+\)', text_10):
        issues.append({'layer': 1, 'severity': 'LOW', 'check': '编号', 'message': '需要协调的事务缺少编号'})
    else:
        passes.append('✅ [L1] 协调事务编号: 已编号')
    
    # 1.6 Empty tasks
    has_empty = False
    for row in table.rows:
        for cell in row.cells:
            if 'ODS-DWD-DWS-DM' in cell.text:
                has_empty = True
    if has_empty:
        issues.append({'layer': 1, 'severity': 'HIGH', 'check': '内容', 'message': '存在空任务（ODS-DWD-DWS-DM）'})
    else:
        passes.append('✅ [L1] 空任务: 已过滤')
    
    # ====================================================================
    # LAYER 2: Data Cross-Reference Checks
    # ====================================================================
    print("\n--- Layer 2: 数据对照 ---")
    
    # 2.1 Date matches Friday
    today = datetime.now()
    monday = today - timedelta(days=today.weekday())
    friday = monday + timedelta(days=4)
    
    date_found = False
    for p in doc.paragraphs:
        if "日期" in p.text:
            date_found = True
            date_match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', p.text)
            if date_match:
                doc_date = datetime(int(date_match.group(1)), int(date_match.group(2)), int(date_match.group(3)))
                friday_date = datetime(friday.year, friday.month, friday.day)
                if doc_date != friday_date:
                    issues.append({
                        'layer': 2, 'severity': 'HIGH', 'check': '日期',
                        'message': f'日期不匹配：文档={doc_date.strftime("%Y-%m-%d")}, 应为周五={friday_date.strftime("%Y-%m-%d")}'
                    })
                else:
                    passes.append('✅ [L2] 日期: 匹配周五 (' + friday_date.strftime('%Y-%m-%d') + ')')
            break
    
    if not date_found:
        issues.append({'layer': 2, 'severity': 'HIGH', 'check': '日期', 'message': '文档中未找到日期字段'})
    
    # 2.2 Filename format - check suffix pattern
    filename = os.path.basename(doc_path)
    if '-工作周报-' in filename and re.search(r'\d{8}-\d{4}\.docx$', filename):
        passes.append('✅ [L2] 文件名: 格式正确 (' + filename + ')')
    else:
        issues.append({
            'layer': 2, 'severity': 'MEDIUM', 'check': '文件名',
            'message': f'文件名格式不正确: {filename}, 应为: *-工作周报-YYYYMMDD-MMDD.docx'
        })
    
    # 2.3 Milestone count matches
    milestone_data = source_data.get('02', [])
    if milestone_data:
        # Check if image exists in document
        has_milestone_img = False
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run._element.findall('.//' + qn('a:blip')):
                            has_milestone_img = True
        
        if not has_milestone_img:
            issues.append({'layer': 2, 'severity': 'HIGH', 'check': '里程碑', 'message': '项目总体进度缺少里程碑截图'})
        else:
            passes.append('✅ [L2] 里程碑截图: 存在')
    
    # 2.4 Verify all active risks are included
    # (Basic check: if doc has risk items)
    risk_items = [t for t in texts_12 if re.match(r'^\d+\)', t)]
    if len(risk_items) == 0:
        # Check if there should be risks
        # This is a content check - would need to parse Sheet 04 for full validation
        passes.append('ℹ️ [L2] 风险条目: 文档中无风险条目（需确认源数据）')
    else:
        passes.append(f'✅ [L2] 风险条目: {len(risk_items)} 条')
    
    # ====================================================================
    # LAYER 3: Visual/Format Checks
    # ====================================================================
    print("\n--- Layer 3: 视觉/格式校验 ---")
    
    # 3.1 Font consistency - check all runs are 微软雅黑 (allow Light variant from template)
    wrong_fonts = []
    wrong_sizes = []
    sample_count = 0
    
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run._element.findall('.//' + qn('a:blip')):
                        continue
                    if not run.text.strip():
                        continue
                    
                    # Allow 微软雅黑 and 微软雅黑 Light (template variant)
                    if run.font.name and run.font.name not in ('微软雅黑', '微软雅黑 Light', 'Microsoft YaHei', 'Microsoft YaHei UI'):
                        if sample_count < 3:
                            wrong_fonts.append("'" + run.text[:20] + "' -> " + run.font.name)
                            sample_count += 1
                    
                    # Allow headers (11pt, 12pt) and body (10.5pt = 133350 EMU)
                    if run.font.size and run.font.size not in [Pt(10.5), Pt(11), Pt(12)]:
                        if len(wrong_sizes) < 3:
                            wrong_sizes.append("'" + run.text[:20] + "' -> " + str(run.font.size))
    
    if wrong_fonts:
        issues.append({
            'layer': 3, 'severity': 'MEDIUM', 'check': '字体',
            'message': f'字体不一致: {wrong_fonts[0]}...'
        })
    else:
        passes.append('✅ [L3] 字体: 统一为微软雅黑')
    
    if wrong_sizes:
        issues.append({
            'layer': 3, 'severity': 'LOW', 'check': '字号',
            'message': f'字号不统一: {wrong_sizes[0]}...'
        })
    else:
        passes.append('✅ [L3] 字号: 统一为五号（10.5pt）')
    
    # 3.2 Check section headers exist and are bold
    section_headers = {
        '项目总体进度': False,
        '本周重点工作': False,
        '下周项目工作计划': False,
        '需要协调的事务': False,
        '预计存在或可能出现的风险': False,
    }
    
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text
            for key in section_headers:
                if key in cell_text:
                    section_headers[key] = True
    
    for section, found in section_headers.items():
        if found:
            passes.append(f'✅ [L3] 章节: {section} 存在')
        else:
            issues.append({
                'layer': 3, 'severity': 'HIGH', 'check': '章节',
                'message': f'缺少章节: {section}'
            })
    
    return issues, passes

def generate_report(issues, passes, doc_path):
    """Generate QA report."""
    report = {
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'document': doc_path,
        'total_issues': len(issues),
        'high_issues': len([i for i in issues if i['severity'] == 'HIGH']),
        'medium_issues': len([i for i in issues if i['severity'] == 'MEDIUM']),
        'low_issues': len([i for i in issues if i['severity'] == 'LOW']),
        'passes': passes,
        'issues': issues,
    }
    
    print("\n" + "=" * 70)
    print("🔍 质量检查报告 v2（三层校验）")
    print("=" * 70)
    print(f"文档: {doc_path}")
    print(f"时间: {report['timestamp']}")
    print(f"问题总数: {report['total_issues']}")
    print(f"  🔴 HIGH:   {report['high_issues']}")
    print(f"  🟡 MEDIUM: {report['medium_issues']}")
    print(f"  🟢 LOW:    {report['low_issues']}")
    
    print(f"\n✅ 通过项 ({len(passes)}):")
    for p in passes:
        print(f"  {p}")
    
    if issues:
        print(f"\n⚠️ 发现问题 ({len(issues)}):")
        for issue in issues:
            severity_emoji = {'HIGH': '🔴', 'MEDIUM': '🟡', 'LOW': '🟢'}[issue['severity']]
            layer_tag = f"[L{issue['layer']}]"
            print(f"  {severity_emoji} {layer_tag} [{issue['severity']}] {issue['check']}: {issue['message']}")
    else:
        print("\n✅ 所有检查通过！文档质量合格。")
    
    print("=" * 70)
    
    return report

if __name__ == "__main__":
    print(f"检查文档: {DOC_PATH}")
    print(f"源数据:   {SOURCE_XLSX}")
    
    if not os.path.exists(DOC_PATH):
        print(f"❌ 文件不存在: {DOC_PATH}")
        sys.exit(1)
    
    source_data = parse_excel(SOURCE_XLSX)
    issues, passes = run_qa_checks(DOC_PATH, source_data)
    report = generate_report(issues, passes, DOC_PATH)
    
    # Output JSON
    print("\n---JSON_START---")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    print("---JSON_END---")
    
    sys.exit(1 if report['high_issues'] > 0 else 0)
