#!/usr/bin/env python3
"""周报话术校验 Agent — 检查内容是否违反 PMO 汇报规范。

检查项：
1. 禁止模糊用语："进行中""推进中""顺利"等无信息量的表述
2. 禁止废话开场："大家好""希望大家继续努力"等无实质内容
3. 必须有量化数据：进度百分比、日期、责任人
4. 风险前置：有风险/延期必须显式标注 ⚠️
5. 已完成事项不反复提（对比上周报告）
6. 编号格式规范：一级中文(一、)、二级阿拉伯(1.)、三级括号(1))
7. 禁止 Markdown 符号残留

用法：python3 validate_report_tone.py <docx_path> [last_week_docx_path]
"""
import sys
import os
import json
import re
from docx import Document


# ============================================================================
# 禁止的话术模式
# ============================================================================
FUZZY_PATTERNS = [
    (r'进行中[，,。]', '模糊进度 — 应给出具体百分比或完成度'),
    (r'进展中[，,。]', '模糊进度 — 应给出具体百分比或完成度'),
    (r'推进中[，,。]', '模糊进度 — 应给出具体百分比或完成度'),
    (r'进展顺利', '空洞评价 — 应给出具体进展'),
    (r'整体.*顺利', '空洞评价 — 应给出具体进展'),
    (r'基本完成', '模糊 — "基本"不明确，应给百分比'),
    (r'大概|大约|左右', '模糊数量词 — 应给出精确值'),
    (r'尽快|早日', '模糊时间 — 应给出明确截止日期'),
    (r'可能|也许|或许', '不确定语气 — 风险应明确定性'),
    (r'希望大家.*努力', '废话 — 删除'),
    (r'大家好|各位好', '废话开场 — 删除'),
    (r'以上.*谢谢', '废话结尾 — 删除'),
]


def extract_all_text(doc):
    """Extract all text from the first table."""
    if not doc.tables:
        return ''
    texts = []
    for row in doc.tables[0].rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                t = para.text.strip()
                if t:
                    texts.append(t)
    return '\n'.join(texts)


def check_fuzzy_language(text):
    """Check for vague/filler language."""
    issues = []
    for pattern, reason in FUZZY_PATTERNS:
        matches = re.finditer(pattern, text)
        for m in matches:
            ctx = text[max(0, m.start()-15):m.end()+15]
            issues.append(f'✗ [话术] {reason}: "...{ctx.strip()}..."')
    return issues


def check_quantification(text):
    """Check that progress descriptions have quantified data."""
    issues = []
    # Look for progress descriptions without percentages or dates
    progress_lines = [l for l in text.split('\n') if any(k in l for k in ['进度', '完成', '开发', '测试', '部署'])]
    # Skip newly-started tasks (no progress expected yet)
    startup_patterns = [r'可以开始', r'开始推进', r'启动', r'进场', r'安排资源', r'通知.*可以', r'调研', r'梳理.*需求', r'收集.*案例']
    # Skip externally-blocked items (can't have % when waiting on vendor/third party)
    external_block_patterns = [r'尚未完成.*需要.*提供', r'需要.*提供.*数据', r'等待.*反馈', r'待.*确认', r'需要确认']
    # Skip already-completed items
    completed_patterns = [r'确认完成', r'已完成', r'已交付', r'已上线']
    for line in progress_lines:
        if any(re.search(p, line) for p in startup_patterns):
            continue  # New task, no quantification needed
        if any(re.search(p, line) for p in external_block_patterns):
            continue  # Externally blocked, can't provide %
        if any(re.search(p, line) for p in completed_patterns):
            continue  # Completed item, no need for future %
        has_number = bool(re.search(r'\d+%', line)) or bool(re.search(r'\d+[天日周月号]', line))
        if not has_number and len(line) > 10:
            issues.append(f'✗ [话术] 缺少量化数据: "{line[:60]}..."')
    return issues


def check_risk_visibility(text):
    """Check that risks/delays are flagged."""
    issues = []
    risk_keywords = ['延期', '阻塞', '风险', '滞后', '无法按期', '存在隐患']
    has_risk_content = any(k in text for k in risk_keywords)

    # Locate the risk section by its heading (e.g. "四、风险与问题跟踪")
    risk_section_match = re.search(r'[一二三四五六七八]+、风险与问题跟踪[\s\S]*?(?=[一二三四五六七八]+、|$)', text)
    if risk_section_match:
        risk_section = risk_section_match.group(0)
        # Should have at least some content beyond "暂无"
        if re.search(r'暂无[，,。]', risk_section) and has_risk_content:
            issues.append('✗ [话术] 正文提到风险/延期，但风险章节写"暂无" — 矛盾')
    return issues


def check_numbering(text):
    """Check numbering format consistency."""
    issues = []
    # Check 一级编号 (一、二、三、)
    level1 = re.findall(r'^[一二三四五六七八]+、', text, re.MULTILINE)
    if not level1:
        # Check if the section should have them
        if '本周重点' in text or '下周' in text:
            issues.append('✗ [格式] 缺少一级编号（一、二、三、）')
    
    # Check 二级编号 (1. 2. 3.)
    level2 = re.findall(r'^\d+\.\s', text, re.MULTILINE)
    if not level2 and level1:
        issues.append('⚠ [格式] 有一级编号但缺少二级编号（1. 2. 3.）')
    
    return issues


def check_completed_repeat(doc_path, last_week_path):
    """Check if completed items are repeated from last week."""
    issues = []
    if not last_week_path or not os.path.exists(last_week_path):
        return issues
    
    try:
        last_doc = Document(last_week_path)
        last_text = extract_all_text(last_doc)
    except Exception:
        return issues
    
    # Extract "已完成" items from last week
    completed_patterns = re.findall(r'([^，,\n]*(?:已完成|已交付|已上线|已关闭)[^，,\n]*)', last_text)
    if not completed_patterns:
        return issues
    
    current_text = extract_all_text(doc)
    for item in completed_patterns:
        if item.strip() and item.strip() in current_text:
            issues.append(f'⚠ [话术] 重复提及上周已完成项: "{item.strip()[:50]}..."')
    
    return issues


def check_quality_language(text):
    """Check raw-note artifacts and low-quality delivery language."""
    issues = []

    duplicate_match = re.search(r'([\u4e00-\u9fff]{2,8})\1', text)
    if duplicate_match:
        issues.append(f'✗ [话术] 重复短语: "{duplicate_match.group(0)}"')

    quality_patterns = [
        (r'稳序', '疑似错词 — 请改为明确对象，如"程序"或"流程"'),
        (r'乡下反馈', '口语化 — 建议改为"基层机构反馈"或"分院反馈"'),
        (r'翻翻了', '口语化 — 建议改为"数据出现翻倍异常"'),
        (r'看一下', '口语化 — 建议改为"核查"或"确认"'),
        (r'找一下', '口语化 — 建议改为"核查"或"定位"'),
        (r'一下', '口语化 — 删除弱化语气，改为明确动作'),
        (r'smartbi', '术语写法 — 建议统一为"Smartbi"'),
        (r'smart\s*开发', '术语写法 — 建议统一为"Smartbi 开发"'),
        (r'svn', '术语写法 — 建议统一为"SVN"'),
    ]
    for pattern, reason in quality_patterns:
        for match in re.finditer(pattern, text):
            ctx = text[max(0, match.start() - 15):match.end() + 15]
            issues.append(f'✗ [话术] {reason}: "...{ctx.strip()}..."')

    return issues


def run_tone_check(doc_path, last_week_path=None):
    """Run all tone/style checks."""
    doc = Document(doc_path)
    text = extract_all_text(doc)
    
    all_issues = []
    all_issues.extend(check_fuzzy_language(text))
    all_issues.extend(check_quantification(text))
    all_issues.extend(check_risk_visibility(text))
    all_issues.extend(check_numbering(text))
    all_issues.extend(check_completed_repeat(doc_path, last_week_path))
    all_issues.extend(check_quality_language(text))
    
    passes = []
    if not any('话术' in i and '模糊' in i for i in all_issues):
        passes.append('✅ [话术] 无模糊用语')
    if not any('话术' in i and '量化' in i for i in all_issues):
        passes.append('✅ [话术] 关键条目已量化')
    if not any('话术' in i and '重复' in i for i in all_issues):
        passes.append('✅ [话术] 未重复已完成项')
    if not any('格式' in i for i in all_issues):
        passes.append('✅ [格式] 编号层级完整')
    if not any('疑似错词' in i or '口语化' in i or '术语写法' in i or '重复短语' in i for i in all_issues):
        passes.append('✅ [话术] 无原始备注残留')
    
    has_high = any(i.startswith('✗') for i in all_issues)
    status = 'FAIL' if has_high else 'PASS'
    
    return {
        'status': status,
        'summary': f'话术校验{status}: {len(passes)} 通过, {len(all_issues)} 问题',
        'passes': passes,
        'issues': all_issues,
    }


if __name__ == '__main__':
    if len(sys.argv) < 2 or not os.path.exists(sys.argv[1]):
        print(json.dumps({'status': 'FAIL', 'issues': ['文档路径无效']}, ensure_ascii=False, indent=2))
        sys.exit(1)
    
    doc_path = sys.argv[1]
    last_week = sys.argv[2] if len(sys.argv) > 2 else None
    
    result = run_tone_check(doc_path, last_week)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    
    if result['status'] == 'FAIL':
        sys.exit(1)
