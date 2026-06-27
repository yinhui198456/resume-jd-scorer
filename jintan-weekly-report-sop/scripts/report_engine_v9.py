#!/usr/bin/env python3
"""
周报生成引擎 v9 - YAML 配置驱动

架构：
  config.yaml → 定义列名/格式/过滤规则/章节结构
  engine.py   → 通用执行引擎（不绑定任何项目）
  
新项目只需：复制 config.yaml → 修改配置 → 运行引擎
"""
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn
import re
import os
import sys
import json
import yaml

# ============================================================================
# YAML Config Loader (with env var expansion)
# ============================================================================
def load_yaml_config(config_path):
    """Load YAML config with env var expansion."""
    with open(config_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Expand environment variables
    content = re.sub(r'\$\{(\w+)\}', lambda m: os.environ.get(m.group(1), ''), content)
    
    return yaml.safe_load(content)

def load_config_simple(config_path):
    """Load config using PyYAML."""
    return load_yaml_config(config_path)

# ============================================================================
# Excel Parser (config-driven column mapping)
# ============================================================================
class ExcelParser:
    def __init__(self, config):
        self.config = config
        self.columns = config.get('source', {}).get('columns', {})
        self.milestone_columns = config.get('source', {}).get('milestone_columns', {})
    
    def parse_excel_serial(self, val_str):
        if not val_str:
            return None
        try:
            return datetime(1899, 12, 30) + timedelta(days=float(val_str))
        except (ValueError, TypeError):
            return None
    
    def parse(self, filepath):
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Excel file not found: {filepath}")
        
        zf = zipfile.ZipFile(filepath)
        ns = {'ss': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
        
        try:
            ss_root = ET.fromstring(zf.read('xl/sharedStrings.xml'))
        except KeyError:
            return {}
        
        ss = []
        for si in ss_root:
            txt = ''.join([t.text or '' for t in si.findall('.//ss:t', ns) + si.findall('.//t') if t.text])
            ss.append(txt)
        
        data = {}
        sheets_config = self.config.get('source', {}).get('sheets', {})
        sheet_map = {}
        for name, sheet_id in sheets_config.items():
            sheet_map[name] = f'xl/worksheets/sheet{sheet_id}.xml'
        
        for name, path in sheet_map.items():
            if path not in zf.namelist():
                continue
            try:
                root = ET.fromstring(zf.read(path))
            except Exception:
                continue
            
            header_map = {}
            rows = []
            last_unit, last_stage = "", ""
            
            unit_col = self.columns.get('work_unit', '工作单元')
            stage_col = self.columns.get('stage', '任务阶段')
            
            for row in root.iter('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}row'):
                cells = {}
                for c in row.iter('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}c'):
                    ref = c.get('r', '')
                    col = ''.join(filter(str.isalpha, ref))
                    v_node = c.find('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v')
                    if c.get('t') == 's' and v_node is not None:
                        val = ss[int(v_node.text)] if v_node.text else ''
                    elif v_node is not None:
                        val = v_node.text
                    else:
                        val = ''
                    cells[col] = val
                
                vals = list(cells.values())
                is_plan_header = self.columns.get('task', '任务') in vals and self.columns.get('responsible', '负责人') in str(vals)
                is_milestone_header = '里程碑' in str(vals) and self.columns.get('status', '状态') in str(vals)
                is_issue_header = '问题描述' in vals and '问题状态' in vals
                
                if is_plan_header or is_milestone_header or is_issue_header:
                    header_map = cells
                    last_unit, last_stage = "", ""
                elif header_map and any(cells.values()):
                    row_dict = {header_map.get(k): v for k, v in cells.items()}
                    
                    if row_dict.get(unit_col, '').strip():
                        last_unit = row_dict[unit_col].strip()
                    else:
                        row_dict[unit_col] = last_unit
                    
                    if row_dict.get(stage_col, '').strip():
                        last_stage = row_dict[stage_col].strip()
                    else:
                        row_dict[stage_col] = last_stage

                    if any(v.strip() for v in row_dict.values()):
                        rows.append(row_dict)
            
            data[name] = rows
        return data

# ============================================================================
# Report Generator (config-driven)
# ============================================================================
class ReportGenerator:
    def __init__(self, config, data):
        self.config = config
        self.data = data
        self.project = config.get('project', {})
        self.format_config = config.get('format', {})
        self.filter_config = config.get('filter', {})
        self.coordination_config = config.get('coordination', {})
        self.sections_config = config.get('sections', {})
        self.columns = config.get('source', {}).get('columns', {})
        
        # Date calculation
        today = datetime.now()
        self.monday = today - timedelta(days=today.weekday())
        self.friday = self.monday + timedelta(days=4)
        self.sunday = self.monday + timedelta(days=6)  # Full week for notes/checks
        self.next_friday = self.friday + timedelta(days=7)
        self.next_sunday = self.sunday + timedelta(days=7)
        
        # Font settings
        self.font_name = self.format_config.get('font', '微软雅黑')
        self.font_size_body = Pt(self.format_config.get('font_size_body', 10.5))
        self.font_size_level1 = Pt(self.format_config.get('font_size_level1', 11))
        self.font_size_header = Pt(self.format_config.get('font_size_header', 12))
        
        self.chinese_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    
    def _set_run_font(self, run, size=None, bold=None):
        """Set both Latin and East Asian font for a run (CJK fix)."""
        run.font.name = self.font_name
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            from lxml import etree
            rFonts = etree.SubElement(rpr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), self.font_name)
        if size is not None:
            run.font.size = size
        if bold is not None:
            run.font.bold = bold
    
    def _fix_template_fonts(self, doc):
        """Fix CJK eastAsia fonts on template header rows and clear placeholder text."""
        # 1. Set eastAsia on all existing CJK runs in the table (template headers)
        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        if re.search(r'[\u4e00-\u9fff]', run.text):
                            rpr = run._element.get_or_add_rPr()
                            from lxml import etree
                            rFonts = rpr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = etree.SubElement(rpr, qn('w:rFonts'))
                            rFonts.set(qn('w:eastAsia'), self.font_name)
        
        # 2. Clear placeholder text in all table cells that contain boilerplate
        placeholder_patterns = [
            '预计存在', '可能出现', '请在此处填写', '示例：',
            '风险描述', '解决方案', '协调事项',
            '本周项目进展情况', '下周项目工作计划',
            '需要协调的事务', '风险与问题跟踪',
            '项目总体进度', '本周重点工作',
        ]
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if any(p in para.text for p in placeholder_patterns):
                        para.clear()
    
    def is_empty_task(self, row):
        """Check if task is empty based on config rules."""
        empty_rules = self.filter_config.get('empty_task', {})
        conditions = empty_rules.get('conditions', [])
        match_all = empty_rules.get('match', 'all') == 'all'
        
        prog = row.get(self.columns.get('progress', '进度'), '').strip()
        notes = row.get(self.columns.get('notes', '备注'), '').strip()
        status = row.get(self.columns.get('status', '状态'), '').strip()
        
        prog_empty = not prog
        if prog:
            try:
                prog_empty = float(prog) == 0
            except ValueError:
                pass
        
        notes_empty = not notes
        status_match = status == '未开始'
        
        if match_all:
            return prog_empty and notes_empty and status_match
        else:
            return prog_empty or notes_empty or status_match
    
    def clean_notes(self, notes):
        """
        Extract note content, prioritizing updates from the current report week.
        e.g. '0509:医管...' will be prioritized if 05/09 is in this week.
        """
        if not notes:
            return ''
        
        # Split notes by newline or pipe
        lines = [line.strip() for line in notes.replace('\n', '|').split('|') if line.strip()]
        if not lines:
            return ''

        # Regex for MMDD: or MM/DD:
        # Matches patterns like 0509:, 05/09:, 05-09:
        date_pattern = re.compile(r'^(\d{2})[\/\.-]?(\d{2})[:：]\s*(.*)')
        
        relevant_notes = []
        
        for line in lines:
            # Check if line starts with a date
            match = date_pattern.match(line)
            if match:
                mm = int(match.group(1))
                dd = int(match.group(2))
                content = match.group(3).strip()
                
                # Construct date (assume current year)
                try:
                    import datetime
                    note_date = datetime.datetime(datetime.datetime.now().year, mm, dd)
                    
                    # Check if within report week (Monday to Friday)
                    if self.monday.date() <= note_date.date() <= self.friday.date():
                        relevant_notes.append(content)
                    else:
                        # If it's an older date, store it as fallback but lower priority
                        relevant_notes.append(line) # Keep original with date for old notes
                except ValueError:
                    relevant_notes.append(line) # Invalid date, keep as is
            else:
                # No date pattern, treat as general note (fallback)
                relevant_notes.append(line)
        
        # Find the LATEST update within the report week
        latest_note = None
        latest_date = None
        
        for line in lines:
            match = date_pattern.match(line)
            if match:
                mm = int(match.group(1))
                dd = int(match.group(2))
                try:
                    note_date = datetime.datetime(datetime.datetime.now().year, mm, dd)
                    # Check if within report week (Monday to Sunday to include weekend updates)
                    if self.monday.date() <= note_date.date() <= self.sunday.date():
                        # If this is the latest date found so far, keep it
                        if latest_date is None or note_date > latest_date:
                            latest_date = note_date
                            latest_note = match.group(3).strip()
                except ValueError:
                    pass
        
        if latest_note:
            return self._sanitize_note_text(latest_note)
            
        # Fallback: If no specific weekly update found, return the first line (cleaned of old date prefixes if any)
        # Or just the last line? Usually the last line in Excel is the latest update.
        last_line = lines[-1]
        # Clean old date prefixes from fallback
        fallback = re.sub(r'^\d{2}[\/\.-]?\d{2}[:：]\s*', '', last_line)
        return self._sanitize_note_text(fallback.strip())
    
    def _sanitize_note_text(self, text):
        """Rephrase vague progress language in notes for tone compliance."""
        if not text:
            return text
        # Strip date prefixes that weren't caught by clean_notes
        text = re.sub(r'^\d{4,6}[\：:]\s*', '', text)
        # Rephrase vague "进行中" to concrete status
        text = re.sub(r'(\S+?)开发进行中', r'\1开发尚未完成', text)
        text = re.sub(r'(?<!\S)进行中(?!\S)', '尚未完成', text)
        return text
    
    def filter_plans(self, plans, week_end=None, is_next_week=False):
        """Filter plans based on config rules."""
        if is_next_week:
            week_end = self.next_friday
        elif week_end is None:
            week_end = self.friday
        
        exclude_statuses = self.filter_config.get('this_week', {}).get('exclude_statuses', ['完成'])
        exclude_units = self.filter_config.get('this_week', {}).get('exclude_work_units', ['项目管理'])
        
        active = []
        for r in plans:
            status = r.get(self.columns.get('status', '状态'), '').strip()
            
            if status in exclude_statuses:
                continue
            
            act_end = self._parse_date(r.get(self.columns.get('plan_end', '计划完成')))
            if act_end and act_end < self.monday and status == '完成':
                continue
            
            plan_start = self._parse_date(r.get(self.columns.get('plan_start', '计划开始')))
            if plan_start is None or plan_start <= week_end:
                unit = r.get(self.columns.get('work_unit', '工作单元'), '').strip()
                if unit not in exclude_units:
                    active.append(r)
        
        return active
    
    def _parse_date(self, val):
        """Parse Excel serial date."""
        if not val:
            return None
        try:
            return datetime(1899, 12, 30) + timedelta(days=float(val))
        except (ValueError, TypeError):
            return None
    
    def build_progress_items(self, plans, source_type='this_week'):
        """Build structured items for progress/plan sections."""
        items = []
        current_unit = ''
        current_stage = ''
        unit_num = 0
        stage_num = 0
        task_num = 0
        
        for r in plans:
            unit = r.get(self.columns.get('work_unit', '工作单元'), '').strip()
            stage = r.get(self.columns.get('stage', '任务阶段'), '').strip()
            task = r.get(self.columns.get('task', '任务'), '').strip()
            
            if not task or self.is_empty_task(r):
                continue
            
            if unit and unit != current_unit:
                current_unit = unit
                current_stage = ''
                unit_num += 1
                stage_num = 0
                task_num = 0
                items.append({'type': 'unit', 'text': unit, 'num': unit_num})
            
            if stage and stage != current_stage:
                current_stage = stage
                stage_num += 1
                task_num = 0
                items.append({'type': 'stage', 'text': stage, 'num': stage_num})
            
            task_num += 1
            notes = self.clean_notes(r.get(self.columns.get('notes', '备注'), ''))
            prog = r.get(self.columns.get('progress', '进度'), '')
            prog_val = ''
            if prog:
                try:
                    p_str = prog.replace('%', '').strip()
                    v = float(p_str)
                    if '%' in prog:
                        v /= 100.0
                    if 0 < v <= 1:
                        prog_val = str(int(v*100)) + "%"
                except ValueError:
                    pass
            
            desc = task
            if notes:
                desc += "：" + notes
            
            # Add action suffix for next week
            if source_type == 'next_week' and prog:
                try:
                    pv_str = prog.replace('%', '').strip()
                    pv = float(pv_str)
                    if '%' in prog:
                        pv /= 100.0
                    if 0 < pv < 1:
                        desc += "（继续推进）"
                    elif pv >= 1:
                        desc += "（确认完成）"
                    else:
                        desc = task + "：启动准备工作"
                except ValueError:
                    desc = task + "：启动准备工作"
            
            if prog_val and source_type == 'this_week':
                desc += " " + prog_val
            
            items.append({'type': 'task', 'text': desc, 'num': task_num})
        
        return items
    
    def find_coordination_items(self, plans):
        """Identify delayed tasks with external reasons."""
        coord_config = self.coordination_config
        if not coord_config.get('enabled', True):
            return []
        
        external_keywords = coord_config.get('external_keywords', [])
        delayed_config = coord_config.get('delayed', {})
        progress_threshold = delayed_config.get('progress_less_than', 1.0)
        
        coordination_items = []
        for r in plans:
            unit = r.get(self.columns.get('work_unit', '工作单元'), '').strip()
            task = r.get(self.columns.get('task', '任务'), '').strip()
            notes = r.get(self.columns.get('notes', '备注'), '').strip()
            progress = r.get(self.columns.get('progress', '进度'), '')
            plan_end = self._parse_date(r.get(self.columns.get('plan_end', '计划完成')))
            status = r.get(self.columns.get('status', '状态'), '').strip()
            
            if not task:
                continue
            
            prog_val = 0
            try:
                if progress:
                    p_str = progress.replace('%', '').strip()
                    prog_val = float(p_str)
                    if '%' in progress:
                        prog_val /= 100.0
                else:
                    prog_val = 0
            except ValueError:
                pass
            
            # Check if delayed
            is_delayed = False
            if plan_end and plan_end < datetime.now() and prog_val < progress_threshold:
                is_delayed = True
            
            # Check "or" conditions
            or_conditions = delayed_config.get('or', [])
            for cond in or_conditions:
                status_cond = cond.get('status_equals', '')
                prog_cond = cond.get('progress_less_than', 1.0)
                if status == status_cond and prog_val < prog_cond:
                    is_delayed = True
            
            # Check external reason
            has_external = any(kw in notes for kw in external_keywords) if notes else False
            
            if is_delayed and has_external:
                issue_summary = notes.split('\n')[0][:80] if notes else ''
                issue_summary = re.sub(r'^\d+[、\.]\s*', '', issue_summary)
                # Strip date prefixes like "0611：" or "20260611："
                issue_summary = re.sub(r'^\d{4,6}[\：:]\s*', '', issue_summary)
                # Rephrase vague "进行中" to concrete status description
                issue_summary = re.sub(r'(\S+?)开发进行中', r'\1开发尚未完成', issue_summary)
                issue_summary = re.sub(r'进行中', '尚未完成', issue_summary)
                coordination_items.append({
                    'unit': unit,
                    'task': task,
                    'progress': prog_val,
                    'issue': issue_summary,
                })
        
        return coordination_items
    
    def build_issues_grouped(self, issues):
        """Build structured issues grouped by config-defined columns."""
        sections_config = self.sections_config.get('issues', {})
        content_config = sections_config.get('content', [{}])[0]
        group_by = content_config.get('group_by', ['数据应用', '模块'])
        
        issue_col = '问题描述'
        status_col = '问题状态'
        app_col = group_by[0] if len(group_by) > 0 else '数据应用'
        mod_col = group_by[1] if len(group_by) > 1 else '模块'
        
        active_issues = [r for r in issues if r.get(status_col, '').strip() not in ('已修复', '已关闭')]
        
        groups = {}
        for r in active_issues:
            app = r.get(app_col, '').strip() or r.get('B', '').strip()
            mod = r.get(mod_col, '').strip() or r.get('C', '').strip()
            desc = r.get(issue_col, '').strip() or r.get('D', '').strip()
            status = r.get(status_col, '').strip() or r.get('H', '').strip()
            
            if not desc:
                continue
            
            desc_simplified = self.clean_notes(desc)
            if not desc_simplified:
                desc_simplified = desc.split('\n')[0][:60]
            
            key = (app, mod)
            if key not in groups:
                groups[key] = []
            groups[key].append({'desc': desc_simplified, 'status': status})
        
        items = []
        current_app = ''
        app_num = 0
        mod_num = 0
        for (app, mod), issues_list in groups.items():
            if app and app != current_app:
                current_app = app
                app_num += 1
                mod_num = 0  # Reset module counter for each app
                items.append({'type': 'app_group', 'text': app, 'num': app_num})
            
            mod_num += 1
            items.append({'type': 'mod_group', 'text': mod if mod else '其他', 'num': mod_num})
            
            for i, issue in enumerate(issues_list, 1):
                status_tag = "[" + issue['status'] + "]" if issue['status'] else ''
                items.append({'type': 'issue', 'text': (issue['desc'] + " " + status_tag).strip(), 'num': i})
        
        return items
    
    def generate_filename(self):
        """Generate filename based on config pattern."""
        pattern = self.config.get('date', {}).get('filename_pattern', '{project_name}-工作周报-{monday_ymd}-{friday_md}.docx')
        project_name = self.project.get('name', '项目')
        
        monday_ymd = self.monday.strftime('%Y%m%d')
        friday_md = self.friday.strftime('%m%d')
        
        filename = pattern.format(
            project_name=project_name,
            monday_ymd=monday_ymd,
            friday_md=friday_md
        )
        return filename
    
    def get_report_date(self):
        """Get report date based on config."""
        date_type = self.config.get('date', {}).get('report_date', 'friday')
        if date_type == 'friday':
            return self.friday
        elif date_type == 'monday':
            return self.monday
        else:
            return datetime.now()
    
    def write_content_to_cell(self, cell, items, section_type='progress'):
        """Write structured items to a Word cell."""
        cell.text = ""
        
        for i, item in enumerate(items):
            if i == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            
            if item['type'] == 'unit':
                num = item.get('num', 1)
                num_str = self.chinese_nums[num - 1] if num <= 10 else str(num)
                run = p.add_run(num_str + "、" + item['text'])
                self._set_run_font(run, size=self.font_size_level1, bold=True)
                p.paragraph_format.space_before = Pt(6)
            
            elif item['type'] == 'stage':
                num = item.get('num', 1)
                run = p.add_run(str(num) + ". " + item['text'])
                self._set_run_font(run, size=self.font_size_body, bold=True)
                p.paragraph_format.left_indent = Cm(0.5)
            
            elif item['type'] == 'task':
                num = item.get('num', 1)
                run = p.add_run(str(num) + ") " + item['text'])
                self._set_run_font(run, size=self.font_size_body)
                p.paragraph_format.left_indent = Cm(1.0)
            
            elif item['type'] == 'app_group':
                num = item.get('num', 1)
                num_str = self.chinese_nums[num - 1] if num <= 10 else str(num)
                run = p.add_run(num_str + "、" + item['text'])
                self._set_run_font(run, size=self.font_size_level1, bold=True)
                p.paragraph_format.space_before = Pt(6)
            
            elif item['type'] == 'mod_group':
                num = item.get('num', 1)
                run = p.add_run(str(num) + ". " + item['text'] + "：")
                self._set_run_font(run, size=self.font_size_body, bold=True)
                p.paragraph_format.left_indent = Cm(0.5)
            
            elif item['type'] == 'issue':
                num = item.get('num', 1)
                run = p.add_run(str(num) + ") " + item['text'])
                self._set_run_font(run, size=self.font_size_body)
                p.paragraph_format.left_indent = Cm(1.0)
            
            elif item['type'] == 'coordination':
                num = item.get('num', 1)
                run = p.add_run(str(num) + ") " + item['text'])
                self._set_run_font(run, size=self.font_size_body)
                p.paragraph_format.left_indent = Cm(0.5)
    
    def generate(self, template_path, output_dir):
        """Generate the report document."""
        print(f"\n📝 步骤: 生成 Word 文档...")
        
        if not os.path.exists(template_path):
            print(f"  ❌ 模板不存在: {template_path}")
            return None
        
        doc = Document(template_path)
        
        # Fix CJK fonts on template header rows (rows 0-5 not overwritten by engine)
        self._fix_template_fonts(doc)
        
        # Update date
        report_date = self.get_report_date()
        for p in doc.paragraphs:
            if "日期" in p.text:
                p.clear()
                date_str = report_date.strftime('%Y年%-m月%-d日')
                run = p.add_run("日期：" + date_str)
                self._set_run_font(run, size=self.font_size_body)
                break
        
        table = doc.tables[0]
        
        # Get data - use config sheet names as keys
        plans = self.data.get('plan', [])
        issues = self.data.get('issues', [])
        
        # Build items
        this_week_plans = self.filter_plans(plans, is_next_week=False)
        this_week_items = self.build_progress_items(this_week_plans, 'this_week')
        print(f"  本周任务: {len(this_week_items)} 条")
        
        next_week_plans = self.filter_plans(plans, is_next_week=True)
        next_week_items = self.build_progress_items(next_week_plans, 'next_week')
        print(f"  下周任务: {len(next_week_items)} 条")
        
        issues_items = self.build_issues_grouped(issues)
        print(f"  问题跟踪: {len([i for i in issues_items if i['type'] == 'issue'])} 条")
        
        coord_items = self.find_coordination_items(plans)
        print(f"  需协调事务: {len(coord_items)} 条")
        
        # Write sections based on config
        sections = self.sections_config
        
        # This week section
        this_week_section = sections.get('this_week', {})
        if this_week_section.get('enabled', True):
            row_idx = this_week_section.get('table_row', 6)
            
            # Write header to Row 5 (header row above content)
            header_row = table.rows[row_idx - 1]
            for hcell in header_row.cells:
                for para in hcell.paragraphs:
                    para.clear()
            header_p = header_row.cells[0].paragraphs[0]
            run = header_p.add_run("一、项目总体进度")
            self._set_run_font(run, size=self.font_size_header, bold=True)
            
            cell = table.cell(row_idx, 0)
            
            # Clear ALL paragraphs in this cell (may span multiple visual rows due to merging)
            for para in list(cell.paragraphs):
                para.clear()
            
            # Add header "一、项目总体进度"
            if cell.paragraphs:
                header_p = cell.paragraphs[0]
            else:
                header_p = cell.add_paragraph()
            run = header_p.add_run("一、项目总体进度")
            self._set_run_font(run, size=self.font_size_header, bold=True)
            header_p.paragraph_format.space_after = Pt(4)
            
            # Add milestone image inline with header
            img_path = "/tmp/milestone_progress_v9.png"
            if os.path.exists(img_path):
                try:
                    # Add image right after header in same paragraph
                    run = header_p.add_run()
                    run.add_picture(img_path, width=Inches(6.5))
                except Exception as e:
                    print(f"  ️ 图片插入失败: {e}")
            
            # Add "本周重点工作：" subheader
            sub_p = cell.add_paragraph()
            run = sub_p.add_run("本周重点工作：")
            self._set_run_font(run, size=self.font_size_body, bold=True)
            sub_p.paragraph_format.space_before = Pt(6)
            sub_p.paragraph_format.space_after = Pt(4)
            
            # Add task items
            for item in this_week_items:
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                
                if item['type'] == 'unit':
                    num = item.get('num', 1)
                    num_str = self.chinese_nums[num - 1] if num <= 10 else str(num)
                    run = p.add_run(num_str + "、" + item['text'])
                    self._set_run_font(run, size=self.font_size_level1, bold=True)
                    p.paragraph_format.space_before = Pt(6)
                elif item['type'] == 'stage':
                    num = item.get('num', 1)
                    run = p.add_run(str(num) + ". " + item['text'])
                    self._set_run_font(run, size=self.font_size_body, bold=True)
                    p.paragraph_format.left_indent = Cm(0.5)
                elif item['type'] == 'task':
                    num = item.get('num', 1)
                    run = p.add_run(str(num) + ") " + item['text'])
                    self._set_run_font(run, size=self.font_size_body)
                    p.paragraph_format.left_indent = Cm(1.0)
        
        # Next week section
        next_week_section = sections.get('next_week', {})
        if next_week_section.get('enabled', True):
            row_idx = next_week_section.get('table_row', 8)
            
            # Write header to Row 7
            header_row = table.rows[row_idx - 1]
            for hcell in header_row.cells:
                for para in hcell.paragraphs:
                    para.clear()
            header_p = header_row.cells[0].paragraphs[0]
            run = header_p.add_run("二、下周项目工作计划")
            self._set_run_font(run, size=self.font_size_header, bold=True)
            
            cell = table.cell(row_idx, 0)
            cell.text = ""
            self.write_content_to_cell(cell, next_week_items, 'next_week')
        
        # Coordination section
        coord_section = sections.get('coordination', {})
        if coord_section.get('enabled', True):
            row_idx = coord_section.get('table_row', 10)
            
            # Write header to Row 9
            header_row = table.rows[row_idx - 1]
            for hcell in header_row.cells:
                for para in hcell.paragraphs:
                    para.clear()
            header_p = header_row.cells[0].paragraphs[0]
            run = header_p.add_run("三、需要协调的事务")
            self._set_run_font(run, size=self.font_size_header, bold=True)
            
            cell = table.cell(row_idx, 0)
            cell.text = ""
            if coord_items:
                for i, ci in enumerate(coord_items):
                    if i == 0:
                        p = cell.paragraphs[0]
                    else:
                        p = cell.add_paragraph()
                    
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    
                    progress_str = ""
                    if ci['progress'] > 0:
                        progress_str = "（当前进度 " + str(int(ci['progress']*100)) + "%）"
                    
                    text = str(i+1) + ") " + ci['unit'] + " - " + ci['task'] + "：" + ci['issue'] + progress_str
                    
                    run = p.add_run(text)
                    self._set_run_font(run, size=self.font_size_body)
            else:
                cell.text = "暂无。"
        
        # Issues section
        issues_section = sections.get('issues', {})
        if issues_section.get('enabled', True):
            row_idx = issues_section.get('table_row', 12)
            
            # Write header to Row 11
            header_row = table.rows[row_idx - 1]
            for hcell in header_row.cells:
                for para in hcell.paragraphs:
                    para.clear()
            header_p = header_row.cells[0].paragraphs[0]
            run = header_p.add_run("四、风险与问题跟踪")
            self._set_run_font(run, size=self.font_size_header, bold=True)

            cell = table.cell(row_idx, 0)
            if issues_items:
                self.write_content_to_cell(cell, issues_items, 'issues')
            else:
                cell.text = "暂无。"
        
        # NOTE: Do NOT delete empty rows - template structure must be preserved
        # Header rows (5, 7, 9, 11) are cleared but should remain as structure
        
        # Save
        filename = self.generate_filename()
        out_path = os.path.join(output_dir, filename)
        doc.save(out_path)
        print(f"  ✅ 文档已保存: {filename}")
        
        return {
            'doc_path': out_path,
            'filename': filename,
            'date': report_date.strftime('%Y年%-m月%-d日'),
            'this_week_count': len(this_week_items),
            'next_week_count': len(next_week_items),
            'issues_count': len([i for i in issues_items if i['type'] == 'issue']),
            'coordination_count': len(coord_items),
            'coordination_items': coord_items,
        }
    
    def generate_email_body(self, report_info):
        """Generate email body text."""
        project_name = self.project.get('name', '')
        short_name = self.project.get('short_name', '')
        date_str = report_info['date']
        
        email_template = """各位领导、同事好：

现将{short_name}项目本周工作进展汇报如下，请审阅。详细周报请见附件。

📅 报告周期：{monday} ~ {friday}

━━━━━━━━━━━━━━━━━━━━━━
📊 本周重点工作（{this_week_count}条）
━━━━━━━━━━━━━━━━━━━━━━
{this_week_summary}

━━━━━━━━━━━━━━━━━━━━━━
📋 下周工作计划（{next_week_count}条）
━━━━━━━━━━━━━━━━━━━━━━
{next_week_summary}

━━━━━━━━━━━━━━━━━━━━━━
🚨 需要协调的事项（{coordination_count}项）
━━━━━━━━━━━━━━━━━━━━━━
{coordination_summary}

━━━━━━━━━━━━━━━━━━━━━━
⚠️ 风险与问题跟踪（{issues_count}条）
━━━━━━━━━━━━━━━━━━━━━━
{issues_summary}

以上为本周项目进展情况，如有问题请随时沟通。

此致
敬礼！

{project_name} 项目组
{date}"""
        
        # Generate summaries
        plans = self.data.get('plan', [])
        this_week_plans = self.filter_plans(plans, is_next_week=False)
        this_week_items = self.build_progress_items(this_week_plans, 'this_week')
        
        # This week summary (first 5 items)
        this_week_lines = []
        current_unit = ''
        for item in this_week_items[:8]:
            if item['type'] == 'unit':
                this_week_lines.append("■ " + item['text'])
                current_unit = item['text']
            elif item['type'] == 'stage':
                this_week_lines.append("  ● " + item['text'])
            elif item['type'] == 'task' and len(this_week_lines) < 10:
                this_week_lines.append("    - " + item['text'][:60])
        this_week_summary = '\n'.join(this_week_lines) if this_week_lines else "本周无新增任务。"
        
        # Next week summary
        next_week_plans = self.filter_plans(plans, is_next_week=True)
        next_week_items = self.build_progress_items(next_week_plans, 'next_week')
        next_week_lines = []
        for item in next_week_items[:8]:
            if item['type'] == 'unit':
                next_week_lines.append("■ " + item['text'])
            elif item['type'] == 'stage':
                next_week_lines.append("  ● " + item['text'])
            elif item['type'] == 'task' and len(next_week_lines) < 10:
                next_week_lines.append("    - " + item['text'][:60])
        next_week_summary = '\n'.join(next_week_lines) if next_week_lines else "下周无计划任务。"
        
        # Coordination summary
        coord_items = report_info.get('coordination_items', [])
        coord_lines = []
        for ci in coord_items:
            progress_pct = str(int(ci['progress']*100)) + "%"
            progress_str = "（" + progress_pct + "）" if ci['progress'] > 0 else ""
            coord_lines.append("• " + ci['unit'] + " - " + ci['task'] + progress_str + "：" + ci['issue'])
        coordination_summary = '\n'.join(coord_lines) if coord_lines else "暂无需要协调的事项。"
        
        # Issues summary
        issues = self.data.get('issues', [])
        issues_items = self.build_issues_grouped(issues)
        issue_lines = []
        for item in issues_items:
            if item['type'] == 'issue' and len(issue_lines) < 8:
                issue_lines.append("• " + item['text'][:60])
        issues_summary = '\n'.join(issue_lines) if issue_lines else "暂无风险与问题。"
        
        monday_str = self.monday.strftime('%Y-%m-%d')
        friday_str = self.friday.strftime('%Y-%m-%d')
        
        body = email_template.format(
            short_name=short_name,
            project_name=project_name,
            monday=monday_str,
            friday=friday_str,
            date=date_str,
            this_week_count=report_info['this_week_count'],
            this_week_summary=this_week_summary,
            next_week_count=report_info['next_week_count'],
            next_week_summary=next_week_summary,
            coordination_count=report_info['coordination_count'],
            coordination_summary=coordination_summary,
            issues_count=report_info['issues_count'],
            issues_summary=issues_summary,
        )
        
        return body

# ============================================================================
# Main Entry Point
# ============================================================================
def resolve_config_paths(config, config_dir):
    """Resolve relative paths in config relative to config file directory."""
    source = config.get('source', {})
    for key in ['xlsx_path', 'template_path', 'output_dir']:
        if key in source and source[key]:
            path = source[key]
            if not os.path.isabs(path):
                source[key] = os.path.normpath(os.path.join(config_dir, path))
    return config

def main():
    config_path = sys.argv[1] if len(sys.argv) > 1 else os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "config.yaml")
    
    if not os.path.exists(config_path):
        print(f"❌ 配置文件不存在: {config_path}")
        sys.exit(1)
    
    print("=" * 60)
    print("周报生成引擎 v9 - YAML 配置驱动")
    print("=" * 60)
    print(f"配置文件: {config_path}")
    
    # Load config
    config = load_config_simple(config_path)
    config_dir = os.path.dirname(os.path.abspath(config_path))
    config = resolve_config_paths(config, config_dir)
    print(f"项目名称: {config.get('project', {}).get('name', 'Unknown')}")
    
    # Parse Excel
    print(f"\n📊 解析 Excel...")
    parser = ExcelParser(config)
    xlsx_path = config.get('source', {}).get('xlsx_path', '')
    data = parser.parse(xlsx_path)
    print(f"  {config['source']['sheets'].get('plan', '?')}-项目计划: {len(data.get(config['source']['sheets'].get('plan', '01'), []))} 行")
    print(f"  {config['source']['sheets'].get('milestone', '?')}-项目里程碑: {len(data.get(config['source']['sheets'].get('milestone', '02'), []))} 行")
    print(f"  {config['source']['sheets'].get('issues', '?')}-应用问题跟踪表: {len(data.get(config['source']['sheets'].get('issues', '04'), []))} 行")
    
    # Generate report
    generator = ReportGenerator(config, data)
    template_path = config.get('source', {}).get('template_path', '')
    output_dir = config.get('source', {}).get('output_dir', '')
    
    report_info = generator.generate(template_path, output_dir)
    
    if report_info:
        # Generate email body
        email_body = generator.generate_email_body(report_info)
        
        # Save email body
        email_path = os.path.join(output_dir, "email_body_" + report_info['filename'].replace('.docx', '.txt'))
        with open(email_path, 'w', encoding='utf-8') as f:
            f.write(email_body)
        print(f"\n  📧 邮件正文已保存: {email_path}")
        
        # ── QA 校验（自动生成后自动执行）──
        doc_full_path = report_info.get('doc_path', '') or os.path.join(output_dir, report_info['filename'])
        qa_script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'validate_weekly_report.py')
        if os.path.exists(qa_script) and os.path.exists(doc_full_path):
            print(f"\n🔍 自动 QA 校验中...")
            import subprocess
            qa_result = subprocess.run(
                ['python3', qa_script, doc_full_path, xlsx_path],
                capture_output=True, text=True, timeout=60
            )
            print(qa_result.stdout)
            if qa_result.returncode != 0:
                print(f"\n⚠️  QA 校验未通过（exit={qa_result.returncode}），请检查上方报告")
                if qa_result.stderr:
                    print(f"STDERR: {qa_result.stderr}")
            else:
                print(f"\n✅ QA 校验通过")
        
        # ── 话术校验（PMO 风格检查）──
        tone_script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'validate_report_tone.py')
        if os.path.exists(tone_script) and os.path.exists(doc_full_path):
            print(f"\n📝 话术校验中...")
            tone_result = subprocess.run(
                ['python3', tone_script, doc_full_path],
                capture_output=True, text=True, timeout=30
            )
            print(tone_result.stdout)
            if tone_result.returncode != 0:
                print(f"\n⚠️  话术校验未通过，请检查上方报告")
            else:
                print(f"\n✅ 话术校验通过")
        
        # Output JSON for delegate_task
        print("\n---JSON_START---")
        result = {
            'report': report_info,
            'email_path': email_path,
            'config_path': config_path,
        }
        print(json.dumps(result, ensure_ascii=False, indent=2))
        print("---JSON_END---")

if __name__ == "__main__":
    main()
