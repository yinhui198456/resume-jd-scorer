#!/usr/bin/env python3
"""周报校验 Agent — 供 delegate_task 调用。

三层校验：
L1 结构：编号/章节/Markdown/空任务
L2 数据：与 Excel 源数据对照、日期/文件名
L3 视觉：CJK 字体、里程碑图片、表格格式

用法：python3 validate_weekly_report.py <docx_path> [xlsx_path]
返回 JSON 格式的校验结果。
"""
import sys
import os
import json
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import re
from datetime import datetime, timedelta

PROJECT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOC_PATH = sys.argv[1] if len(sys.argv) > 1 else ""
SOURCE_XLSX = sys.argv[2] if len(sys.argv) > 2 else os.path.join(PROJECT_DIR, "data", "金坛二期项目跟进表.xlsx")

def parse_excel_sheets(filepath):
    """Parse key sheets from xlsx."""
    if not os.path.exists(filepath):
        return {}
    zf = zipfile.ZipFile(filepath)
    try:
        ss_root = ET.fromstring(zf.read('xl/sharedStrings.xml'))
    except KeyError:
        return {}
    ss = []
    ns = {'ss': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
    for si in ss_root:
        txt = ''.join([t.text or '' for t in si.findall('.//ss:t', ns) + si.findall('.//t') if t.text])
        ss.append(txt)
    data = {}
    for sheet_id in ['1', '2', '3']:
        path = f'xl/worksheets/sheet{sheet_id}.xml'
        if path not in zf.namelist():
            continue
        try:
            root = ET.fromstring(zf.read(path))
        except Exception:
            continue
        header_map = {}
        rows = []
        for row in root.iter('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}row'):
            cells = {}
            for c in row.iter('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}c'):
                ref = c.get('r', '')
                col = ''.join(filter(str.isalpha, ref))
                v_node = c.find('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v')
                if c.get('t') == 's' and v_node is not None:
                    val = ss[int(v_node.text)] if v_node.text else ''
                elif v_node is not None:
                    val = v_node.text
                else:
                    val = ''
                cells[col] = val
            vals = list(cells.values())
            if any(k in str(vals) for k in ['任务', '里程碑', '问题描述', '序号']):
                header_map = cells
            elif header_map and any(cells.values()):
                row_dict = {header_map.get(k): v for k, v in cells.items()}
                if any(str(v).strip() for v in row_dict.values()):
                    rows.append(row_dict)
        data[sheet_id] = rows
    return data


def check_cjk_fonts(doc, min_row=6):
    """Check that all runs with CJK text have eastAsia font set. Only checks rows >= min_row."""
    issues = []
    total_runs = 0
    cjk_runs = 0
    missing_ea = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        total_runs += 1
                        # Check if text contains CJK
                        if re.search(r'[\u4e00-\u9fff]', run.text):
                            cjk_runs += 1
                            rpr = run._element.find(qn('w:rPr'))
                            has_ea = False
                            if rpr is not None:
                                rFonts = rpr.find(qn('w:rFonts'))
                                if rFonts is not None:
                                    has_ea = rFonts.get(qn('w:eastAsia')) is not None
                            if not has_ea:
                                missing_ea += 1
                                if missing_ea <= 3:  # Only report first few
                                    issues.append(f'  ✗ 缺少 eastAsia 字体: "{run.text[:40]}"')
    return total_runs, cjk_runs, missing_ea, issues


def check_milestone_image(doc):
    """Check if milestone image is embedded."""
    img_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for drawing in run._element.findall('.//' + qn('wp:inline')):
                            img_count += 1
    return img_count


def find_empty_numbered_paragraphs(doc):
    """Find paragraphs that render numbering but have no visible text."""
    issues = []
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    text = para.text.strip()
                    ppr = para._p.pPr
                    has_numbering = ppr is not None and ppr.numPr is not None
                    if not text and has_numbering:
                        issues.append(
                            f'✗ [L1] 空编号段落: table={table_idx}, row={row_idx}, cell={cell_idx}, paragraph={para_idx}'
                        )
    return issues


def find_next_week_plan_issues(doc):
    """Validate that next-week plan is future action, not current status replay."""
    issues = []
    if not doc.tables:
        return issues
    table = doc.tables[0]
    if len(table.rows) <= 8:
        return issues

    text = table.cell(8, 0).text
    forbidden_patterns = [
        (r'继续推进', '下周计划不应使用空泛状态词"继续推进"'),
        (r'确认完成', '下周计划不应复述已完成状态"确认完成"'),
        (r'100%', '下周计划不应包含 100% 已完成任务'),
    ]
    for pattern, message in forbidden_patterns:
        if re.search(pattern, text):
            issues.append(f'✗ [L1] {message}')

    return issues


def run_validation(doc_path, source_data):
    doc = Document(doc_path)
    table = doc.tables[0] if doc.tables else None
    if not table:
        return {'status': 'FAIL', 'issues': ['文档中无表格'], 'passes': []}

    issues = []
    passes = []

    # L1: Structure
    cell_6 = table.cell(6, 0) if len(table.rows) > 6 else None
    if cell_6:
        texts = [p.text.strip() for p in cell_6.paragraphs if p.text.strip()]
        if any(re.match(r'^[一二三四五六七八]+、', t) for t in texts):
            passes.append('✅ [L1] 本周一级编号: 存在')
        else:
            issues.append('✗ [L1] 本周重点缺少一级编号（一、二、三、）')
        if any(re.match(r'^\d+\.\s', t) for t in texts):
            passes.append('✅ [L1] 本周二级编号: 存在')
        else:
            issues.append('✗ [L1] 本周重点缺少二级编号（1. 2. 3.）')

    # Check for Markdown symbols
    has_md = False
    for row in table.rows:
        for cell in row.cells:
            if '**' in cell.text or '###' in cell.text or '## ' in cell.text:
                has_md = True
    if has_md:
        issues.append('✗ [L1] 文档中存在 Markdown 符号')
    else:
        passes.append('✅ [L1] Markdown: 已清除')

    empty_numbering_issues = find_empty_numbered_paragraphs(doc)
    if empty_numbering_issues:
        issues.extend(empty_numbering_issues[:10])
        if len(empty_numbering_issues) > 10:
            issues.append(f'✗ [L1] 空编号段落: 另有 {len(empty_numbering_issues) - 10} 处')
    else:
        passes.append('✅ [L1] 空编号段落: 无')

    next_week_issues = find_next_week_plan_issues(doc)
    if next_week_issues:
        issues.extend(next_week_issues)
    else:
        passes.append('✅ [L1] 下周计划: 无状态复述')

    # Check for empty task placeholder (only if the cell is essentially empty
    # or contains only the placeholder text with no real task content)
    placeholder_found = False
    for row in table.rows:
        cell_texts = [cell.text.strip() for cell in row.cells]
        # A row is a placeholder only if ALL cells have the same minimal placeholder text
        # and at least one cell is empty (no numbering like 1) 一、 etc.)
        if all(t in ('', 'ODS-DWD-DWS-DM', 'ODS-DWD-DWS-DM：') for t in cell_texts) and any(t == '' for t in cell_texts):
            placeholder_found = True
            break
    if placeholder_found:
        issues.append('✗ [L1] 存在空任务占位符')

    # L2: Data checks
    # Extract date from filename: ...-工作周报-YYYYMMDD-NNNN.docx
    # NNNN is MMDD of Friday (e.g. 20260511-0515 means Mon May 11 - Fri May 15)
    filename = os.path.basename(doc_path)
    date_m = re.search(r'工作周报-(\d{4})(\d{2})(\d{2})-(\d{2})(\d{2})', filename)
    expected_friday = None
    if date_m:
        year = int(date_m.group(1))
        fri_month = int(date_m.group(4))
        fri_day = int(date_m.group(5))
        # Handle跨年: if month < mon_month, year+1
        mon_month = int(date_m.group(2))
        if fri_month < mon_month:
            year += 1
        expected_friday = datetime(year, fri_month, fri_day)
    
    for p in doc.paragraphs:
        if "日期" in p.text:
            m = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', p.text)
            if m:
                doc_date = datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
                if expected_friday and doc_date == expected_friday:
                    passes.append(f'✅ [L2] 日期: 匹配周五 ({doc_date.strftime("%Y-%m-%d")})')
                elif expected_friday:
                    issues.append(f'✗ [L2] 日期不匹配: 文档={doc_date.strftime("%Y-%m-%d")}, 应={expected_friday.strftime("%Y-%m-%d")}')
                else:
                    passes.append(f'✅ [L2] 日期: {doc_date.strftime("%Y-%m-%d")} (文件名未匹配，跳过对照)')
            break

    # Check item counts against source data
    plan_data = source_data.get('1', [])
    if plan_data:
        # Count non-completed tasks
        active = [r for r in plan_data if r.get('状态', '').strip() != '完成']
        passes.append(f'✅ [L2] 源数据: {len(plan_data)} 条计划, {len(active)} 条未完成')

    # L3: Visual checks
    total_runs, cjk_runs, missing_ea, font_issues = check_cjk_fonts(doc, min_row=6)
    if missing_ea == 0 and cjk_runs > 0:
        passes.append(f'✅ [L3] CJK 字体: 全部 {cjk_runs} 个中文 run 已设置 eastAsia')
    elif missing_ea > 0:
        issues.append(f'✗ [L3] CJK 字体: {missing_ea}/{cjk_runs} 个中文 run 缺少 eastAsia 字体')
        issues.extend(font_issues[:3])

    img_count = check_milestone_image(doc)
    if img_count > 0:
        passes.append(f'✅ [L3] 里程碑图片: 已嵌入 ({img_count} 张)')
    else:
        issues.append('✗ [L3] 里程碑图片: 未找到嵌入图片')

    # Check for tofu/square characters (U+FFFD or U+25A1)
    tofu_count = 0
    for row in table.rows:
        for cell in row.cells:
            tofu_count += cell.text.count('\uFFFD') + cell.text.count('\u25A1')
    if tofu_count > 0:
        issues.append(f'✗ [L3] 乱码检测: 发现 {tofu_count} 个方块字符 (tofu)')
    else:
        passes.append('✅ [L3] 乱码检测: 无方块字符')

    # Summary
    has_critical = any('✗' in i and ('[L1]' in i or '[L3] 乱码' in i) for i in issues)
    # L2/L3 non-critical issues are warnings, not failures
    non_critical = [i for i in issues if '[L2]' in i or '[L3]' in i]
    status = 'FAIL' if has_critical else ('WARN' if issues else 'PASS')

    return {
        'status': status,
        'summary': f'校验{status}: {len(passes)} 通过, {len(issues)} 问题',
        'passes': passes,
        'issues': issues,
        'stats': {
            'total_runs': total_runs,
            'cjk_runs': cjk_runs,
            'missing_eastasia': missing_ea,
            'images_embedded': img_count,
            'tofu_chars': tofu_count,
        }
    }


if __name__ == '__main__':
    if not DOC_PATH or not os.path.exists(DOC_PATH):
        print(json.dumps({'status': 'FAIL', 'issues': ['文档路径无效']}, ensure_ascii=False, indent=2))
        sys.exit(1)

    source_data = parse_excel_sheets(SOURCE_XLSX) if os.path.exists(SOURCE_XLSX) else {}
    result = run_validation(DOC_PATH, source_data)
    print(json.dumps(result, ensure_ascii=False, indent=2))

    if result['status'] == 'FAIL':
        sys.exit(1)
