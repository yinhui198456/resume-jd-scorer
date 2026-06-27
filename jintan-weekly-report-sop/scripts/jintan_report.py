#!/usr/bin/env python3
"""
Jintan Phase II Weekly Report Generator (Word Version) - v3
Improvements:
1. Structured Output: Replaced pipe-separated text with hierarchical paragraphs.
2. Result Check: Added a verification stage to print the structure before saving.
3. Clean Formatting: Uses bolding for Units/Stages and bullets for tasks.
"""
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm, RGBColor
import re
import os

# Configuration
PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(PROJECT_DIR, "..", "templates", "常州市金坛第一人民医院数据指挥中心二期项目-工作周报-20260420-0424.docx")
SOURCE_XLSX = os.path.join(PROJECT_DIR, "..", "data", "金坛二期项目跟进表.xlsx")
OUTPUT_DIR = os.path.join(PROJECT_DIR, "..", "output")

TODAY = datetime.now()
MONDAY = TODAY - timedelta(days=TODAY.weekday())
SUNDAY = MONDAY + timedelta(days=6)

def parse_excel_serial(val_str):
    if not val_str: return None
    try: return datetime(1899, 12, 30) + timedelta(days=float(val_str))
    except: return None

def parse_excel(filepath):
    if not os.path.exists(filepath): return {}
    zf = zipfile.ZipFile(filepath)
    ns = {'ss': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
    
    try:
        ss_root = ET.fromstring(zf.read('xl/sharedStrings.xml'))
    except: return {}
    ss = []
    for si in ss_root:
        txt = ''.join([t.text or '' for t in si.findall('.//ss:t', ns) + si.findall('.//t') if t.text])
        ss.append(txt)
        
    data = {}
    # Mapping: Sheet Name -> XML Path
    sheet_map = {'01': 'xl/worksheets/sheet2.xml', '04': 'xl/worksheets/sheet5.xml'}
    
    for name, path in sheet_map.items():
        if path not in zf.namelist(): continue
        try:
            root = ET.fromstring(zf.read(path))
        except: continue
        header_map, rows = {}, []
        last_unit, last_stage = "", ""
        
        for row in root.findall('.//ss:row', ns):
            cells = {}
            for c in row.findall('.//ss:c', ns):
                ref, col = c.get('r'), ''.join(filter(str.isalpha, c.get('r')))
                v_node = c.find('ss:v', ns)
                val = ss[int(v_node.text)] if c.get('t') == 's' and v_node is not None else (v_node.text if v_node is not None else '')
                cells[col] = val
            
            vals = list(cells.values())
            is_plan_header = "任务" in vals and ("负责人" in str(vals) or "负责人-内" in vals) and "状态" in vals
            is_issue_header = "问题描述" in vals and "问题状态" in vals
            
            if is_plan_header or is_issue_header:
                header_map = cells
                last_unit, last_stage = "", ""
            elif header_map and any(cells.values()):
                row_dict = {header_map.get(k): v for k, v in cells.items()}
                
                if row_dict.get('工作单元', '').strip():
                    last_unit = row_dict['工作单元'].strip()
                else:
                    row_dict['工作单元'] = last_unit
                    
                if row_dict.get('任务阶段', '').strip():
                    last_stage = row_dict['任务阶段'].strip()
                else:
                    row_dict['任务阶段'] = last_stage

                if any(row_dict.values()):
                    rows.append(row_dict)
        data[name] = rows
    return data

def build_structured_list(plans):
    """Converts plan rows into a structured list for reporting."""
    structured_items = []
    last_unit = ""
    last_stage = ""
    
    # Ensure we keep original order or sort if necessary. 
    # Original Excel order is usually Unit -> Stage -> Task.
    # We rely on 'last_unit' logic to detect hierarchy changes.
    
    for r in plans:
        unit = r.get('工作单元', '').strip()
        stage = r.get('任务阶段', '').strip()
        task = r.get('任务', '').strip()
        
        if not task: continue
        
        # Clean notes
        notes = r.get('备注', '').strip().replace('\n', ' ')
        prog = r.get('进度', '')
        prog_val = ""
        if prog:
            try:
                v = float(prog)
                if 0 < v < 1: prog_val = f"{int(v*100)}%"
            except: pass
            
        desc = task
        if notes:
            # Remove prefixes like "1、xxx："
            clean_notes = re.sub(r'^\d+[、\.]\s*\S+：\s*', '', notes)
            if clean_notes:
                 desc += f"：{clean_notes}"
        if prog_val:
            desc += f" ({prog_val})"
            
        # Build Structure
        if unit and unit != last_unit:
            structured_items.append({'type': 'unit', 'text': unit})
            last_unit = unit
            last_stage = "" # Reset stage
            
        if stage and stage != last_stage:
            structured_items.append({'type': 'stage', 'text': stage})
            last_stage = stage
            
        structured_items.append({'type': 'task', 'text': desc})
        
    return structured_items

def check_result(structured_items):
    """Stage: Result Check - Prints the structure to console."""
    print("\n=== 阶段：结果检查 (Structure Check) ===")
    print(f"Total items generated: {len(structured_items)}")
    if not structured_items:
        print("⚠️ WARNING: No data found! The report might be empty.")
        return
        
    print("Preview of structure:")
    for item in structured_items[:15]: # Limit preview
        if item['type'] == 'unit':
            print(f"📂 [UNIT] {item['text']}")
        elif item['type'] == 'stage':
            print(f"   📁 [STAGE] {item['text']}")
        else:
            print(f"      - {item['text'][:60]}{'...' if len(item['text'])>60 else ''}")
    if len(structured_items) > 15:
        print("... (truncated)")
    print("==============================================\n")

def generate_report():
    print("--- Generating Jintan Weekly Report (v3 - Structured) ---")
    data = parse_excel(SOURCE_XLSX)
    
    # 1. Filter Plans (01-Project Plan)
    plans = data.get('01', [])
    active_plans = []
    for r in plans:
        status = r.get('状态', '').strip()
        act_end = parse_excel_serial(r.get('实际完成', ''))
        
        # Exclude historical completed
        if act_end and act_end < MONDAY and status == '完成':
            continue
        
        # Include active/upcoming
        plan_start = parse_excel_serial(r.get('计划开始', ''))
        if plan_start is None or plan_start <= SUNDAY:
            active_plans.append(r)
            
    # Exclude "Project Management" from progress
    active_plans = [r for r in active_plans if r.get('工作单元', '').strip() != '项目管理']

    # 2. Filter Issues (04-Application Problem Tracking)
    issues = data.get('04', [])
    # Exclude "已修复"
    active_issues = [r for r in issues if r.get('问题状态', '').strip() != '已修复']

    # 3. Build Structure
    progress_items = build_structured_list(active_plans)
    
    # 4. Check Result (New Stage)
    print(">>> Checking This Week's Progress...")
    check_result(progress_items)
    
    # 5. Write to Word
    if not os.path.exists(TEMPLATE_PATH):
        print("Error: Template not found.")
        return
    
    doc = Document(TEMPLATE_PATH)

    # Update Date
    for p in doc.paragraphs:
        if "日期" in p.text:
            p.clear()
            run = p.add_run(f"日期：{TODAY.strftime('%Y年%m月%d日')}")
            run.font.size = Pt(12)
            run.font.bold = True
            break

    table = doc.tables[0]
    
    def write_structured_content(cell, items):
        """Writes a list of structured items to a Word cell with formatting."""
        cell.text = ""
        for i, item in enumerate(items):
            # Add paragraph
            if i == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            
            # Reset spacing
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            run = p.add_run(item['text'])
            run.font.size = Pt(10) # Base size
            run.font.name = '宋体' # Standard font
            
            if item['type'] == 'unit':
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
                p.paragraph_format.space_before = Pt(6)
            elif item['type'] == 'stage':
                run.font.bold = True
                p.paragraph_format.left_indent = Cm(0.3)
            else: # task
                p.paragraph_format.left_indent = Cm(0.7)
                # Prepend bullet
                run.text = "• " + item['text']

    if progress_items:
        write_structured_content(table.cell(6, 1), progress_items)
        write_structured_content(table.cell(8, 1), progress_items) # Same for Next Week for now
        
    # Issues as simple text list (or structured if preferred)
    issue_text_parts = []
    for r in active_issues:
        mod = r.get('模块', '').strip()
        desc = r.get('问题描述', '').strip().replace('\n', ' ')
        if desc:
            issue_text_parts.append(f"{mod}：{desc}")
            
    if issue_text_parts:
        # Simple paragraph for risks
        cell = table.cell(12, 1)
        cell.text = ""
        for i, txt in enumerate(issue_text_parts):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            run = p.add_run(txt)
            run.font.size = Pt(10)
    else:
        table.cell(12, 1).text = "暂无。"

    out_name = f"常州市金坛第一人民医院数据指挥中心二期项目-工作周报-{TODAY.strftime('%Y%m%d')}.docx"
    out_path = os.path.join(OUTPUT_DIR, out_name)
    doc.save(out_path)
    print(f"✅ Report saved to: {out_path}")

if __name__ == "__main__":
    generate_report()
