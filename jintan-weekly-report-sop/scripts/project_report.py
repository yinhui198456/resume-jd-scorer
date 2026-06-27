#!/usr/bin/env python3
"""
Project Weekly Report Generator (v5).
Generates a formal Word report with:
1. Summary from Milestones (Sheet 02).
2. Structured Tasks (Sheet 01) with Semantic Rewriting (SOW Context).
3. Risk/Issue Tracking (Sheet 04).
"""
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

# ============================================================
# QA Fix Functions (v5.1) — TDD implemented
# ============================================================

_CHINESE_NUMS = "零一二三四五六七八九"

def _to_chinese_num(n):
    """阿拉伯数字 → 中文数字（1-10）"""
    if 1 <= n <= 10:
        return _CHINESE_NUMS[n]
    return str(n)


def sanitize_note_text(text):
    """
    清理备注文本：
    1. "进行中" → "尚未完成"
    2. 剥离日期前缀 (如 "0611：")
    3. 多日期备注提取最新的（按日期逆序）
    """
    if not text:
        return ""

    # 多日期备注：提取最新的 (格式: "0428:xxx|0611:yyy" 或 "0428:\nxxx\n0611:\nyyy")
    # 先尝试按 | 分割
    if "|" in text:
        parts = text.split("|")
        text = parts[-1]
    else:
        # 按日期前缀分割 (如 "0424:\nxxx\n0509:\nyyy")
        date_parts = re.findall(r'(\d{4})[:：]\s*(.*?)(?=\d{4}[:：]|$)', text, re.DOTALL)
        if date_parts:
            # 取最后一个（最新日期）
            text = date_parts[-1][1].strip()

    # 剥离日期前缀 (如 "0611：" 或 "0611:")
    text = re.sub(r'^\d{4}[:：]\s*', '', text)
    text = re.sub(r'^\d{2}[:：]\s*', '', text)

    # "进行中" → "尚未完成"
    text = text.replace("进行中", "尚未完成")

    return text.strip()


def is_external_blocked(text):
    """
    判断是否为第三方阻塞项（白名单模式）
    匹配以下模式则返回 True：
    - 尚未完成...需要提供...
    - 需要...提供...数据
    - 等待...反馈
    """
    if not text:
        return False

    patterns = [
        r'尚未完成.*需要.*提供',
        r'需要.*提供.*数据',
        r'等待.*反馈',
        r'待.*确认',
        r'需.*协调',
    ]

    return any(re.search(p, text) for p in patterns)


def clean_placeholder_text(text):
    """
    清理模板占位文字
    移除：预计存在、可能出现、解决方案 等占位符
    """
    if not text:
        return ""

    placeholder_patterns = [
        r'预计存在或可能出现的风险及解决方案',
        r'预计存在',
        r'可能出现的风险',
        r'暂无。',
    ]

    result = text
    for pattern in placeholder_patterns:
        result = re.sub(re.escape(pattern), '', result)

    return result.strip()


def format_with_numbering(items):
    """
    为条目添加编号：
    - unit 类型：一、二、三、（中文数字）
    - stage 类型：1. 2. 3.（阿拉伯数字）
    - task 类型：保持原样（前面已有 • ）
    """
    if not items:
        return []

    result = []
    unit_counter = 0
    stage_counter = 0

    for item in items:
        item_type = item.get('type', '')
        text = item.get('text', '')

        if item_type == 'unit':
            unit_counter += 1
            stage_counter = 0
            new_text = f"{_to_chinese_num(unit_counter)}、{text}"
            result.append({'type': 'unit', 'text': new_text})

        elif item_type == 'stage':
            stage_counter += 1
            new_text = f"{stage_counter}. {text}"
            result.append({'type': 'stage', 'text': new_text})

        else:
            result.append(item)

    return result


# ============================================================
# End QA Fix Functions
# ============================================================

# Configuration
PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(PROJECT_DIR, "..", "templates", "常州市金坛第一人民医院数据指挥中心二期项目-工作周报-20260420-0424.docx")
SOURCE_XLSX = os.path.join(PROJECT_DIR, "..", "data", "金坛二期项目跟进表.xlsx")
OUTPUT_DIR = os.path.join(PROJECT_DIR, "..", "output")

TODAY = datetime.now()
MONDAY = TODAY - timedelta(days=TODAY.weekday())
SUNDAY = MONDAY + timedelta(days=6)

# Context from Feasibility Report (SOW)
# This allows the SOP to "understand" technical terms like ODS/DWD
SOW_CONTEXT = {
    "ODS": "数据贴源层 (ODS) 数据接入",
    "DWD": "数据明细层 (DWD) 清洗建模",
    "DWS": "数据汇总层 (DWS) 指标构建",
    "DM": "数据应用层 (DM) 场景开发",
    "湖仓一体": "基于湖仓一体架构的数据底座建设"
}

def parse_excel_serial(val_str):
    if not val_str: return None
    try: return datetime(1899, 12, 30) + timedelta(days=float(val_str))
    except: return None

def parse_excel(filepath):
    if not os.path.exists(filepath): return {}
    zf = zipfile.ZipFile(filepath)
    ns = {'ss': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
    
    try: ss_root = ET.fromstring(zf.read('xl/sharedStrings.xml'))
    except: return {}
    ss = []
    for si in ss_root:
        txt = ''.join([t.text or '' for t in si.findall('.//ss:t', ns) + si.findall('.//t') if t.text])
        ss.append(txt)
        
    data = {}
    wb_root = ET.fromstring(zf.read('xl/workbook.xml'))
    sheet_map = {}
    
    # Map sheet relationships
    for s in wb_root.findall('.//ss:sheet', ns):
        name = s.get('name')
        rid = s.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        try:
            rels_root = ET.fromstring(zf.read('xl/_rels/workbook.xml.rels'))
            for rel in rels_root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                if rel.get('Id') == rid:
                    sheet_map[name] = 'xl/' + rel.get('Target')
                    break
        except: pass

    # Parse specific sheets
    for name, path in sheet_map.items():
        if '00-待办' in name: data['00'] = _parse_sheet(zf, path, ss, ns, type='plan')
        elif '01-项目计划' in name: data['01'] = _parse_sheet(zf, path, ss, ns, type='plan')
        elif '02-项目里程碑' in name: data['02'] = _parse_sheet(zf, path, ss, ns, type='milestone')
        elif '04-应用问题' in name: data['04'] = _parse_sheet(zf, path, ss, ns, type='issue')
        
    return data

def _parse_sheet(zf, path, ss, ns, type='plan'):
    try: root = ET.fromstring(zf.read(path))
    except: return []
    
    header_map, rows = {}, []
    last_unit = ""
    
    for row in root.findall('.//ss:row', ns):
        cells = {}
        for c in row.findall('.//ss:c', ns):
            ref = c.get('r')
            col = ''.join(filter(str.isalpha, ref))
            v_node = c.find('ss:v', ns)
            val = ss[int(v_node.text)] if c.get('t') == 's' and v_node is not None else (v_node.text if v_node is not None else '')
            cells[col] = val
        
        vals = list(cells.values())
        is_header = False
        if type == 'plan': is_header = "任务" in vals and "负责人" in str(vals)
        elif type == 'milestone': is_header = "里程碑名称" in vals
        elif type == 'issue': is_header = "问题描述" in vals and "问题状态" in vals
        
        if is_header:
            header_map = cells
            if type != 'plan': last_unit = ""
        elif header_map and any(cells.values()):
            # QA Fix: Map by column letter to header name (handles missing columns)
            row_dict = {}
            for col_letter, col_name in header_map.items():
                row_dict[col_name] = cells.get(col_letter, '')
            
            if type == 'plan':
                if row_dict.get('工作单元', '').strip(): last_unit = row_dict['工作单元'].strip()
                else: row_dict['工作单元'] = last_unit
                
            if any(row_dict.values()):
                rows.append(row_dict)
    return rows

def generate_summary(data):
    milestones = data.get('02', [])
    # Find active milestone — 状态"进行中"不改写（这是里程碑状态，非备注）
    active = [m for m in milestones if m.get('状态', '').strip() == '进行中']
    if not active: return "项目暂无进行中的里程碑。"

    m = active[0]
    parts = []
    parts.append(f"当前项目处于 **【{m.get('里程碑名称', '')}】** 阶段。")
    if m.get('里程碑标志'): parts.append(f"核心目标：{m['里程碑标志']}。")
    # Check deliverables — 清理占位文字
    dels = m.get('交付件', '').replace('\n', '、')
    dels = clean_placeholder_text(dels)
    if dels and len(dels) < 150: parts.append(f"关键交付件：{dels}。")
    return " ".join(parts)

def generate_tasks(data, include_unstarted=False):
    """
    生成任务列表。
    include_unstarted=False: 只列出有进度的任务（本周事项）
    include_unstarted=True: 列出所有任务（下周计划）
    """
    plans = data.get('01', [])
    active_plans = []

    for r in plans:
        status = r.get('状态', '').strip()
        act_end = parse_excel_serial(r.get('实际完成', ''))

        # Filter historical completed
        if act_end and act_end < MONDAY and status == '完成': continue

        # Filter future
        plan_start = parse_excel_serial(r.get('计划开始', ''))
        if plan_start is None or plan_start <= SUNDAY:
            active_plans.append(r)

    # Remove Project Management
    active_plans = [r for r in active_plans if r.get('工作单元', '').strip() != '项目管理']

    # Structure
    items = []
    last_unit = ""
    last_stage = ""

    for r in active_plans:
        unit = r.get('工作单元', '').strip()
        stage = r.get('任务阶段', '').strip()
        task = r.get('任务', '').strip()
        notes = r.get('备注', '').strip()
        prog = r.get('进度', '')

        # Parse progress
        prog_val = 0
        if prog:
            try:
                prog_str = str(prog).strip().replace('%', '')
                prog_val = float(prog_str)
                if prog_val > 1: prog_val = prog_val  # Already percentage
                else: prog_val = prog_val * 100  # Decimal to percentage
            except: prog_val = 0

        # QA Fix: Skip unstarted tasks for this-week report
        if not include_unstarted and prog_val <= 0:
            continue

        # Semantic Rewriting
        if task in SOW_CONTEXT and not notes:
            if prog_val > 0:
                task_desc = SOW_CONTEXT[task]
                task_desc += f" ({int(prog_val)}%)"
            else:
                continue
        else:
            # Standard Task
            if notes: notes = re.sub(r'^\d+[、\.]\s*\S+：\s*', '', notes)
            # QA Fix: sanitize note text (进行中→尚未完成, strip date prefix, latest only)
            notes = sanitize_note_text(notes)
            task_desc = task
            if notes: task_desc += f"：{notes}"

            # Show progress percentage
            if prog_val > 0:
                task_desc += f" ({int(prog_val)}%)"

            # Skip empty tasks
            if task in SOW_CONTEXT and not notes and prog_val <= 0:
                continue

        if not task_desc: continue

        # Hierarchy
        if unit and unit != last_unit:
            items.append({'type': 'unit', 'text': unit})
            last_unit = unit
            last_stage = ""

        if stage and stage != last_stage:
            items.append({'type': 'stage', 'text': stage})
            last_stage = stage

        items.append({'type': 'task', 'text': f"• {task_desc}"})

    # QA Fix: Apply Chinese/decimal numbering
    return format_with_numbering(items)


def generate_next_week_tasks(data):
    """生成下周计划 — 列出计划开始在下周的任务，不包含问题"""
    plans = data.get('01', [])
    next_week_plans = []

    for r in plans:
        plan_start = parse_excel_serial(r.get('计划开始', ''))
        status = r.get('状态', '').strip()

        # Skip completed
        if status == '完成': continue

        # Include tasks starting next week or later but not too far
        if plan_start and plan_start > SUNDAY:
            next_friday = SUNDAY + timedelta(days=7)
            if plan_start <= next_friday:
                next_week_plans.append(r)

    # Structure
    items = []
    last_unit = ""
    last_stage = ""

    for r in next_week_plans:
        unit = r.get('工作单元', '').strip()
        stage = r.get('任务阶段', '').strip()
        task = r.get('任务', '').strip()

        if unit and unit != last_unit:
            items.append({'type': 'unit', 'text': unit})
            last_unit = unit
            last_stage = ""

        if stage and stage != last_stage:
            items.append({'type': 'stage', 'text': stage})
            last_stage = stage

        items.append({'type': 'task', 'text': f"• {task}"})

    return format_with_numbering(items)


def generate_risks(data):
    issues = data.get('04', [])
    active = [r for r in issues if r.get('问题状态', '').strip() != '已修复']

    risks = []
    for r in active:
        mod = r.get('模块', '').strip()
        desc = r.get('问题描述', '').strip().replace('\n', ' ')
        # QA Fix: clean placeholder text
        desc = clean_placeholder_text(desc)
        if desc: risks.append(f"• {mod}：{desc[:50]}{'...' if len(desc)>50 else ''}")
    return risks


def _clean_template_placeholders(doc):
    """
    QA Fix: 清理模板中的占位文字
    扫描所有表格单元格，移除"预计存在或可能出现的风险及解决方案"等占位符
    """
    placeholder_patterns = [
        '预计存在或可能出现的风险及解决方案',
        '预计存在',
        '可能出现的风险',
        '暂无。',
    ]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text
                    cleaned = original
                    for placeholder in placeholder_patterns:
                        cleaned = cleaned.replace(placeholder, '')
                    if cleaned.strip() != original.strip():
                        # Replace paragraph content
                        para.clear()
                        if cleaned.strip():
                            run = para.add_run(cleaned.strip())
                            run.font.size = Pt(10)
                            run.font.name = '宋体'


def write_to_doc(summary_text, milestone_data, this_week_tasks, next_week_tasks, risk_items):
    if not os.path.exists(TEMPLATE_PATH): return
    doc = Document(TEMPLATE_PATH)

    # QA Fix: Clean placeholder text from template cells
    _clean_template_placeholders(doc)

    # Date
    for p in doc.paragraphs:
        if "日期" in p.text:
            p.clear()
            run = p.add_run(f"日期：{TODAY.strftime('%Y年%m月%d日')}")
            run.font.size = Pt(12)
            break

    table = doc.tables[0]
    
    def write_cell(cell, summary, milestone_data, this_week_tasks, next_week_tasks, risks, is_milestone_cell=False):
        cell.text = ""
        
        # Handle milestone cell specially
        if is_milestone_cell:
            generate_milestone_table(milestone_data, cell, summary)
            return
        
        # 1. Summary with milestone table
        p_title = cell.paragraphs[0]
        run_title = p_title.add_run("一、项目总体进度")
        run_title.bold = True; run_title.font.size = Pt(12); run_title.font.name = '黑体'
        
        p_sum = cell.add_paragraph()
        p_sum.paragraph_format.space_after = Pt(4)
        run_sum = p_sum.add_run(summary)
        run_sum.font.size = Pt(10); run_sum.font.name = '宋体'
        
        # 2. This week tasks (only with progress)
        if this_week_tasks:
            p2 = cell.add_paragraph()
            run2 = p2.add_run("二、本周重点工作\n")
            run2.bold = True; run2.font.size = Pt(12); run2.font.name = '黑体'
            
            for item in this_week_tasks:
                p_item = cell.add_paragraph()
                p_item.paragraph_format.space_before = Pt(1)
                p_item.paragraph_format.space_after = Pt(1)
                
                run_i = p_item.add_run(item['text'])
                run_i.font.size = Pt(10); run_i.font.name = '宋体'
                
                if item['type'] == 'unit':
                    run_i.bold = True; run_i.font.size = Pt(11); run_i.font.name = '黑体'
                    p_item.paragraph_format.space_before = Pt(4)
                elif item['type'] == 'stage':
                    run_i.bold = True; run_i.font.size = Pt(10)
                    p_item.paragraph_format.left_indent = Cm(0.5)
                else:
                    p_item.paragraph_format.left_indent = Cm(1.0)
        
        # 3. Next week plan (no risks)
        if next_week_tasks:
            p3 = cell.add_paragraph()
            run3 = p3.add_run("三、下周计划\n")
            run3.bold = True; run3.font.size = Pt(12); run3.font.name = '黑体'
            
            for item in next_week_tasks:
                p_item = cell.add_paragraph()
                p_item.paragraph_format.space_before = Pt(1)
                p_item.paragraph_format.space_after = Pt(1)
                
                run_i = p_item.add_run(item['text'])
                run_i.font.size = Pt(10); run_i.font.name = '宋体'
                
                if item['type'] == 'unit':
                    run_i.bold = True; run_i.font.size = Pt(11); run_i.font.name = '黑体'
                    p_item.paragraph_format.space_before = Pt(4)
                elif item['type'] == 'stage':
                    run_i.bold = True; run_i.font.size = Pt(10)
                    p_item.paragraph_format.left_indent = Cm(0.5)
                else:
                    p_item.paragraph_format.left_indent = Cm(1.0)
                    
        # 4. Risks (separate section)
        if risks:
            # Clear any template placeholder text first
            cell.text = ""
            p4 = cell.add_paragraph()
            run4 = p4.add_run("四、风险与问题跟踪\n")
            run4.bold = True; run4.font.size = Pt(12); run4.font.name = '黑体'
            
            for r_text in risks:
                p_r = cell.add_paragraph()
                run_r = p_r.add_run(r_text)
                run_r.font.size = Pt(10); run_r.font.name = '宋体'
                p_r.paragraph_format.left_indent = Cm(0.5)

    # Row 5: 项目总体进度（独立第一章）
    write_cell(table.cell(5, 1), summary_text, milestone_data, [], [], [], is_milestone_cell=True)
    # Row 6: 本周重点工作
    write_cell(table.cell(6, 1), "", "", this_week_tasks, [], [])
    # Row 8: 下周计划
    write_cell(table.cell(8, 1), "", "", [], next_week_tasks, [])
    # Row 11: 风险与问题跟踪（独立章节）
    write_cell(table.cell(11, 1), "", "", [], [], risk_items)
    
    out_name = f"常州市金坛第一人民医院数据指挥中心二期项目-工作周报-{TODAY.strftime('%Y%m%d')}.docx"
    doc.save(os.path.join(OUTPUT_DIR, out_name))

def generate_milestone_table(milestone_data, target_cell, summary_text):
    """生成里程碑表格（Word 表格对象）"""
    # Clear cell first
    target_cell.text = ""
    
    # Add title
    p_title = target_cell.paragraphs[0]
    run_title = p_title.add_run("一、项目总体进度")
    run_title.bold = True
    run_title.font.size = Pt(12)
    run_title.font.name = '黑体'
    
    # Add summary
    p_sum = target_cell.add_paragraph()
    p_sum.paragraph_format.space_after = Pt(6)
    run_sum = p_sum.add_run(summary_text)
    run_sum.font.size = Pt(10)
    run_sum.font.name = '宋体'
    
    # Get active milestones
    active_milestones = [m for m in milestone_data if m.get('里程碑名称', '')]
    
    # Create milestone table directly in the cell
    table = target_cell.add_table(rows=len(active_milestones[:5]) + 1, cols=4)
    # table.style = 'Table Grid'  # Template doesn't have this style
    
    # Header row
    headers = ['编号', '里程碑名称', '状态', '计划完成']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '黑体'
    
    # Data rows
    for idx, m in enumerate(active_milestones[:5]):
        row_idx = idx + 1
        table.rows[row_idx].cells[0].text = m.get('编号', str(idx+1))
        table.rows[row_idx].cells[1].text = m.get('里程碑名称', '')
        table.rows[row_idx].cells[2].text = m.get('状态', '')
        table.rows[row_idx].cells[3].text = m.get('里程碑时点', '')
        
        for cell in table.rows[row_idx].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = '宋体'
    
    return target_cell




def quality_check_report(doc_path):
    """质量检查：验证生成的周报是否符合规范"""
    issues = []
    passes = []
    
    doc = Document(doc_path)
    table = doc.tables[0]
    
    # Check 1: 项目总体进度章节存在
    row5_text = table.rows[5].cells[1].text
    if '项目总体进度' in row5_text:
        passes.append('✅ 项目总体进度章节存在')
    else:
        issues.append('❌ 缺少项目总体进度章节')
    
    # Check 2: 里程碑表格存在
    row5_cell = table.rows[5].cells[1]
    has_milestone_table = len(row5_cell.tables) > 0
    has_milestone_text = '里程碑' in row5_text or '已完成' in row5_text or '进行中' in row5_text
    if has_milestone_table or has_milestone_text:
        passes.append('✅ 里程碑表格存在')
    else:
        issues.append('❌ 缺少里程碑表格')
    
    # Check 3: 本周重点工作章节存在
    row6_text = table.rows[6].cells[1].text
    if '本周重点工作' in row6_text:
        passes.append('✅ 本周重点工作章节存在')
    else:
        issues.append('❌ 缺少本周重点工作章节')
    
    # Check 4: 风险章节独立
    row11_text = table.rows[11].cells[1].text
    if '风险与问题跟踪' in row11_text:
        passes.append('✅ 风险章节独立')
    else:
        issues.append('❌ 风险章节未独立')
    
    # Check 5: 风险未混入本周任务
    if '风险' not in row6_text and '问题跟踪' not in row6_text:
        passes.append('✅ 风险未混入本周任务')
    else:
        issues.append('❌ 风险混入本周任务')
    
    # Check 6: 编号格式
    if '一、' in row6_text and '1.' in row6_text:
        passes.append('✅ 编号格式正确（一级中文、二级阿拉伯）')
    else:
        issues.append(' 编号格式不正确')
    
    return {
        'status': 'PASS' if not issues else 'FAIL',
        'passes': passes,
        'issues': issues
    }

def main():
    print("=== Generating v5 SOP Report ===")
    data = parse_excel(SOURCE_XLSX)

    print("\n[1] 检查项目里程碑 (Sheet 02)...")
    summary = generate_summary(data)
    milestones = data.get('02', [])
    print(f"    -> {summary}")

    print("\n[2] 检查本周重点任务 (有进度)...")
    this_week_tasks = generate_tasks(data, include_unstarted=False)
    print(f"    -> 共提取 {len(this_week_tasks)} 个有进度条目")
    for t in this_week_tasks[:5]: print(f"       {t['text'][:50]}...")

    print("\n[3] 检查下周计划 (未启动)...")
    next_week_tasks = generate_next_week_tasks(data)
    print(f"    -> 共提取 {len(next_week_tasks)} 个下周条目")
    for t in next_week_tasks[:5]: print(f"       {t['text'][:50]}...")

    print("\n[4] 检查风险问题 (Sheet 04)...")
    risks = generate_risks(data)
    print(f"    -> 共提取 {len(risks)} 个未修复问题")

    print("\n[5] 生成 Word 报告...")
    out_name = f"常州市金坛第一人民医院数据指挥中心二期项目-工作周报-{TODAY.strftime('%Y%m%d')}.docx"
    write_to_doc(summary, milestones, this_week_tasks, next_week_tasks, risks)
    print("✅ Done.")

    print("\n[6] 质量检查...")
    out_path = os.path.join(OUTPUT_DIR, out_name)
    qc_result = quality_check_report(out_path)
    print(f"    状态: {qc_result['status']}")
    for p in qc_result['passes']: print(f"    {p}")
    for i in qc_result['issues']: print(f"    {i}")

if __name__ == "__main__":
    main()